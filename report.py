"""
QUANT REPORT GENERATOR v4 — Pure Python DOCX
Pipeline: quant_pipeline.py → quant_report.xlsx → report.py → .docx

Usage: python report.py [excel_path] [output_path] [top_n]
  Default: python report.py quant_report.xlsx bao_cao_quant.docx 5

Dependencies: pip install pandas openpyxl python-docx
"""
import sys, logging, os, io
from datetime import datetime
import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
log = logging.getLogger("report")

# ═══════════════════════════════════════════════════════════════
# COLORS
# ═══════════════════════════════════════════════════════════════
CLR_BLUE = RGBColor(0x1F, 0x4E, 0x79)
CLR_ACCENT = RGBColor(0x2E, 0x75, 0xB6)
CLR_DARK = RGBColor(0x33, 0x33, 0x33)
CLR_GRAY = RGBColor(0x66, 0x66, 0x66)
CLR_RED = RGBColor(0xC0, 0x00, 0x00)
CLR_GREEN = RGBColor(0x2E, 0x7D, 0x32)
CLR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# ═══════════════════════════════════════════════════════════════
# CHART FACTORY (matplotlib → BytesIO for embedding in docx)
# ═══════════════════════════════════════════════════════════════
CHART_STYLE = {
    'figure.facecolor': 'white', 'axes.facecolor': '#FAFAFA',
    'axes.edgecolor': '#CCCCCC', 'axes.labelcolor': '#333333',
    'text.color': '#333333', 'xtick.color': '#666666', 'ytick.color': '#666666',
    'grid.color': '#EEEEEE', 'grid.alpha': 0.8, 'font.size': 9,
    'font.family': 'sans-serif',
}

class Charts:
    DPI = 150

    @staticmethod
    def _to_buf(fig):
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=Charts.DPI, bbox_inches='tight',
                    facecolor=fig.get_facecolor(), edgecolor='none')
        plt.close(fig); buf.seek(0)
        return buf

    @classmethod
    def score_ranking(cls, symbols, scores):
        """Horizontal bar chart — Top 5 Quant Score."""
        plt.rcParams.update(CHART_STYLE)
        fig, ax = plt.subplots(figsize=(7, 2.8))
        y = np.arange(len(symbols))
        colors = ['#1B5E20' if s>=80 else '#2E7D32' if s>=65 else '#F57F17' if s>=50 else '#E65100' if s>=35 else '#B71C1C' for s in scores]
        bars = ax.barh(y, scores, color=colors, edgecolor='white', height=0.6, alpha=0.9)
        for bar, sc in zip(bars, scores):
            ax.text(bar.get_width()+1.2, bar.get_y()+bar.get_height()/2,
                    str(int(sc)), va='center', ha='left', fontsize=11, fontweight='bold', color='#333')
        ax.set_yticks(y); ax.set_yticklabels(symbols, fontsize=11, fontweight='bold')
        ax.set_xlim(0, 105); ax.set_xlabel('Quant Score'); ax.invert_yaxis()
        for t, l, c in [(80,'SWING BUY','#1B5E20'),(65,'BUY','#2E7D32'),(50,'WATCH','#F57F17')]:
            ax.axvline(t, color=c, linewidth=0.8, linestyle='--', alpha=0.5)
            ax.text(t, -0.6, l, fontsize=7, color=c, ha='center')
        ax.grid(axis='x', alpha=0.3)
        ax.set_title('Bảng xếp hạng Quant Score', fontsize=12, fontweight='bold', color='#1F4E79', pad=10)
        fig.tight_layout()
        return cls._to_buf(fig)

    @classmethod
    def risk_reward_map(cls, symbols, entries, stops, tp2s, tp3s):
        """Risk/Reward bar chart."""
        plt.rcParams.update(CHART_STYLE)
        fig, ax = plt.subplots(figsize=(7, 2.8))
        y = np.arange(len(symbols))
        for i in range(len(symbols)):
            e, sl, t2, t3 = entries[i], stops[i], tp2s[i], tp3s[i]
            if e <= 0: continue
            rp = (e-sl)/e*100 if sl>0 else 0
            r2 = (t2-e)/e*100 if t2>0 else 0
            r3 = (t3-e)/e*100 if t3>0 else 0
            ax.barh(y[i], -rp, left=0, height=0.5, color='#C62828', alpha=0.8)
            ax.barh(y[i], r2, left=0, height=0.5, color='#2E7D32', alpha=0.7)
            if r3 > r2:
                ax.barh(y[i], r3-r2, left=r2, height=0.5, color='#66BB6A', alpha=0.5)
            if rp > 0:
                ax.text(-rp/2, y[i], f'-{rp:.1f}%', ha='center', va='center', fontsize=7, color='white', fontweight='bold')
            if r3 > 0:
                ax.text(r3+0.5, y[i], f'+{r3:.1f}%', ha='left', va='center', fontsize=7, color='#2E7D32')
        ax.set_yticks(y); ax.set_yticklabels(symbols, fontsize=10, fontweight='bold')
        ax.axvline(0, color='#333', linewidth=1, alpha=0.5)
        ax.set_xlabel('% so với Entry'); ax.invert_yaxis()
        ax.grid(axis='x', alpha=0.2)
        ax.set_title('Bản đồ Rủi ro / Lợi nhuận', fontsize=12, fontweight='bold', color='#1F4E79', pad=10)
        fig.tight_layout()
        return cls._to_buf(fig)

    @classmethod
    def performance_comparison(cls, symbols, sharpes, win_rates, max_dds):
        """Grouped bar chart comparing Sharpe, WinRate, MaxDD across stocks."""
        plt.rcParams.update(CHART_STYLE)
        fig, axes = plt.subplots(1, 3, figsize=(9, 3))

        colors = ['#1F4E79', '#2E75B6', '#5B9BD5', '#BDD7EE', '#9DC3E6']
        x = np.arange(len(symbols))

        # Sharpe
        ax = axes[0]
        bars = ax.bar(x, sharpes, color=colors[:len(symbols)], width=0.6, edgecolor='white')
        ax.set_xticks(x); ax.set_xticklabels(symbols, fontsize=8, fontweight='bold')
        ax.set_title('Sharpe Ratio', fontsize=10, fontweight='bold', color='#1F4E79')
        ax.axhline(1.0, color='#2E7D32', linewidth=0.8, linestyle='--', alpha=0.6)
        ax.grid(axis='y', alpha=0.3)
        for bar, v in zip(bars, sharpes):
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.05, f'{v:.2f}',
                    ha='center', va='bottom', fontsize=8, fontweight='bold')

        # WinRate
        ax = axes[1]
        bars = ax.bar(x, win_rates, color=colors[:len(symbols)], width=0.6, edgecolor='white')
        ax.set_xticks(x); ax.set_xticklabels(symbols, fontsize=8, fontweight='bold')
        ax.set_title('Win Rate (%)', fontsize=10, fontweight='bold', color='#1F4E79')
        ax.axhline(50, color='#F57F17', linewidth=0.8, linestyle='--', alpha=0.6)
        ax.set_ylim(0, 100); ax.grid(axis='y', alpha=0.3)
        for bar, v in zip(bars, win_rates):
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+1, f'{v:.0f}%',
                    ha='center', va='bottom', fontsize=8, fontweight='bold')

        # MaxDD (negative values, show as positive for visual)
        ax = axes[2]
        dd_abs = [abs(d) for d in max_dds]
        bars = ax.bar(x, dd_abs, color=['#C62828' if d>20 else '#F57F17' if d>15 else '#2E7D32' for d in dd_abs],
                      width=0.6, edgecolor='white')
        ax.set_xticks(x); ax.set_xticklabels(symbols, fontsize=8, fontweight='bold')
        ax.set_title('Max Drawdown (%)', fontsize=10, fontweight='bold', color='#1F4E79')
        ax.axhline(20, color='#C62828', linewidth=0.8, linestyle='--', alpha=0.6)
        ax.grid(axis='y', alpha=0.3)
        for bar, v in zip(bars, max_dds):
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.5, f'{v:.1f}%',
                    ha='center', va='bottom', fontsize=7, fontweight='bold')

        fig.suptitle('So sánh hiệu suất Top 5', fontsize=12, fontweight='bold', color='#1F4E79', y=1.02)
        fig.tight_layout()
        return cls._to_buf(fig)

    @classmethod
    def risk_heatmap(cls, symbols, ploss3s, lock_dds, confs):
        """Risk metrics comparison — horizontal grouped bars."""
        plt.rcParams.update(CHART_STYLE)
        fig, axes = plt.subplots(1, 3, figsize=(9, 2.8))

        x = np.arange(len(symbols))

        # PLoss3%
        ax = axes[0]
        colors_pl = ['#C62828' if p>30 else '#F57F17' if p>20 else '#2E7D32' for p in ploss3s]
        bars = ax.bar(x, ploss3s, color=colors_pl, width=0.6, edgecolor='white')
        ax.set_xticks(x); ax.set_xticklabels(symbols, fontsize=8, fontweight='bold')
        ax.set_title('PLoss3% (Rủi ro T+3)', fontsize=10, fontweight='bold', color='#1F4E79')
        ax.axhline(20, color='#F57F17', linewidth=0.8, linestyle='--', alpha=0.6, label='Ngưỡng 20%')
        ax.grid(axis='y', alpha=0.3)
        for bar, v in zip(bars, ploss3s):
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.5, f'{v:.1f}%',
                    ha='center', va='bottom', fontsize=7, fontweight='bold')

        # LockDD%
        ax = axes[1]
        ld_abs = [abs(d) for d in lock_dds]
        colors_ld = ['#C62828' if d>8 else '#F57F17' if d>5 else '#2E7D32' for d in ld_abs]
        bars = ax.bar(x, ld_abs, color=colors_ld, width=0.6, edgecolor='white')
        ax.set_xticks(x); ax.set_xticklabels(symbols, fontsize=8, fontweight='bold')
        ax.set_title('Lock DD% (DD trong T+3)', fontsize=10, fontweight='bold', color='#1F4E79')
        ax.grid(axis='y', alpha=0.3)
        for bar, v in zip(bars, lock_dds):
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.2, f'{v:.1f}%',
                    ha='center', va='bottom', fontsize=7, fontweight='bold')

        # Confidence
        ax = axes[2]
        colors_cf = ['#2E7D32' if c>=67 else '#F57F17' if c>=50 else '#C62828' for c in confs]
        bars = ax.bar(x, confs, color=colors_cf, width=0.6, edgecolor='white')
        ax.set_xticks(x); ax.set_xticklabels(symbols, fontsize=8, fontweight='bold')
        ax.set_title('Confidence (%)', fontsize=10, fontweight='bold', color='#1F4E79')
        ax.axhline(67, color='#2E7D32', linewidth=0.8, linestyle='--', alpha=0.6)
        ax.set_ylim(0, 100); ax.grid(axis='y', alpha=0.3)
        for bar, v in zip(bars, confs):
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+1, f'{v:.0f}%',
                    ha='center', va='bottom', fontsize=8, fontweight='bold')

        fig.suptitle('So sánh rủi ro Top 5', fontsize=12, fontweight='bold', color='#1F4E79', y=1.02)
        fig.tight_layout()
        return cls._to_buf(fig)

    @classmethod
    def stock_radar(cls, symbol, factors):
        """Radar chart for individual stock — 5 normalized factors."""
        plt.rcParams.update(CHART_STYLE)
        labels = list(factors.keys())
        values = list(factors.values())
        n = len(labels)
        if n < 3: return None

        angles = np.linspace(0, 2*np.pi, n, endpoint=False).tolist()
        values_p = values + [values[0]]
        angles += [angles[0]]

        fig, ax = plt.subplots(figsize=(3.5, 3.5), subplot_kw=dict(polar=True))
        ax.fill(angles, [max(0,v) for v in values_p], color='#2E75B6', alpha=0.2)
        ax.plot(angles, [max(0,v) for v in values_p], color='#2E75B6', linewidth=2)
        for a, v in zip(angles[:-1], values):
            c = '#2E7D32' if v > 0.5 else '#F57F17' if v > 0 else '#C62828'
            ax.plot(a, max(0,v), 'o', color=c, markersize=7, zorder=5)
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(labels, fontsize=8, color='#333')
        ax.set_ylim(0, 1.0)
        ax.set_yticks([0.25, 0.5, 0.75])
        ax.set_yticklabels(['','',''], fontsize=5)
        ax.grid(color='#CCC', alpha=0.5)
        ax.set_title(f'{symbol} — Radar chỉ số', fontsize=11, fontweight='bold', color='#1F4E79', pad=15)
        fig.tight_layout()
        return cls._to_buf(fig)

    @classmethod
    def price_forecast(cls, symbol, dates, actual, predicted,
                       upper=None, lower=None, forecast_start_idx=None):
        """Price Forecast chart — Actual vs Predicted with optional confidence band.
        
        Args:
            symbol: Stock ticker
            dates: list/array of datetime or str dates
            actual: list/array of actual close prices
            predicted: list/array of predicted prices (same length as actual)
            upper: optional upper confidence band
            lower: optional lower confidence band
            forecast_start_idx: index where out-of-sample forecast begins (vertical line)
        """
        plt.rcParams.update(CHART_STYLE)
        fig, ax = plt.subplots(figsize=(8, 3.5))

        x = np.arange(len(dates))

        # ── Plot actual price ──
        ax.plot(x, actual, color='#1F4E79', linewidth=1.8, label='Giá thực tế',
                marker='o', markersize=2.5, zorder=3)

        # ── Plot predicted price ──
        ax.plot(x, predicted, color='#2E7D32', linewidth=1.8, label='Dự báo',
                marker='o', markersize=2.5, zorder=3)

        # ── Confidence band ──
        if upper is not None and lower is not None:
            ax.fill_between(x, lower, upper, color='#2E7D32', alpha=0.1,
                            label='Khoảng tin cậy 95%')

        # ── Forecast start vertical line ──
        if forecast_start_idx is not None and 0 < forecast_start_idx < len(dates):
            ax.axvline(forecast_start_idx, color='#C62828', linewidth=1,
                       linestyle='--', alpha=0.7, zorder=2)
            ax.text(forecast_start_idx + 0.3, ax.get_ylim()[1] * 0.98,
                    'Bắt đầu dự báo', fontsize=7, color='#C62828',
                    va='top', ha='left', fontstyle='italic')

        # ── Formatting ──
        # X-axis: show ~8-10 date labels max
        n_ticks = min(10, len(dates))
        tick_step = max(1, len(dates) // n_ticks)
        tick_positions = list(range(0, len(dates), tick_step))
        if (len(dates) - 1) not in tick_positions:
            tick_positions.append(len(dates) - 1)
        ax.set_xticks(tick_positions)
        # Format date labels
        date_labels = []
        for i in tick_positions:
            d = dates[i]
            if hasattr(d, 'strftime'):
                date_labels.append(d.strftime('%d/%m'))
            else:
                date_labels.append(str(d)[-5:])
        ax.set_xticklabels(date_labels, fontsize=7, rotation=30, ha='right')

        ax.set_ylabel('Giá (VND)', fontsize=9)
        ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'{v:,.0f}'))
        ax.legend(loc='upper left', fontsize=8, framealpha=0.9)
        ax.grid(axis='both', alpha=0.3)
        ax.set_title(f'{symbol} — Giá thực tế vs Dự báo',
                     fontsize=12, fontweight='bold', color='#1F4E79', pad=10)

        # ── Metrics annotation box ──
        # Calculate MAPE and Directional Accuracy on overlapping non-NaN region
        mask = np.array([not (np.isnan(a) or np.isnan(p))
                         for a, p in zip(actual, predicted)])
        if mask.sum() > 1:
            a_clean = np.array(actual)[mask]
            p_clean = np.array(predicted)[mask]
            mape = np.mean(np.abs((a_clean - p_clean) / a_clean)) * 100
            # Directional accuracy
            if len(a_clean) > 1:
                a_dir = np.diff(a_clean)
                p_dir = np.diff(p_clean)
                da = np.mean(np.sign(a_dir) == np.sign(p_dir)) * 100
            else:
                da = 0
            metrics_text = f'MAPE: {mape:.1f}%  |  DA: {da:.0f}%'
            ax.text(0.98, 0.02, metrics_text, transform=ax.transAxes,
                    fontsize=8, ha='right', va='bottom',
                    bbox=dict(boxstyle='round,pad=0.3', facecolor='white',
                              edgecolor='#CCCCCC', alpha=0.9))

        fig.tight_layout()
        return cls._to_buf(fig)

    @classmethod
    def ml_forecast(cls, symbol, dates, actual, xgb_pred=None, lstm_pred=None,
                    ens_pred=None, ci_upper=None, ci_lower=None, forecast_start_idx=None):
        """Multi-model ML Forecast chart — XGBoost, LSTM, Ensemble with CI band.

        Args:
            symbol: Stock ticker
            dates: list of datetime or str dates
            actual: list of actual close prices (full series incl. forecast zone)
            xgb_pred: XGBoost predictions (from forecast_start_idx onward, or full-length with NaN)
            lstm_pred: LSTM predictions (same convention)
            ens_pred: Ensemble predictions (same convention)
            ci_upper/ci_lower: confidence interval for ensemble
            forecast_start_idx: index where out-of-sample forecast begins
        """
        plt.rcParams.update(CHART_STYLE)
        fig, ax = plt.subplots(figsize=(8, 3.8))
        n = len(dates)
        x = np.arange(n)

        # ── Helper: pad short forecast arrays to full length with NaN ──
        def _pad(arr):
            if arr is None:
                return None
            arr = np.array(arr, dtype=float)
            if len(arr) == n:
                return arr
            if forecast_start_idx is not None and len(arr) == n - forecast_start_idx:
                return np.concatenate([np.full(forecast_start_idx, np.nan), arr])
            return arr  # already full or unknown alignment

        xgb = _pad(xgb_pred)
        lstm = _pad(lstm_pred)
        ens = _pad(ens_pred)
        ci_up = _pad(ci_upper)
        ci_lo = _pad(ci_lower)

        # ── Actual price ──
        ax.plot(x, actual, color='#1F4E79', linewidth=1.8, label='Giá thực tế',
                marker='o', markersize=2, zorder=4)

        # ── Model predictions ──
        model_lines = [
            (xgb,  '#378ADD', 'XGBoost',  (5, 3)),
            (lstm, '#D85A30', 'LSTM',     (5, 3)),
            (ens,  '#1D9E75', 'Ensemble', None),
        ]
        for pred, color, label, dash in model_lines:
            if pred is not None and not np.all(np.isnan(pred)):
                style = dict(color=color, linewidth=2 if dash is None else 1.5,
                             label=label, zorder=3, marker='o', markersize=2)
                if dash:
                    style['linestyle'] = '--'
                    style['dashes'] = dash
                ax.plot(x, pred, **style)

        # ── Confidence band ──
        if ci_up is not None and ci_lo is not None:
            ax.fill_between(x, ci_lo, ci_up, color='#1D9E75', alpha=0.1,
                            label='95% CI (Ensemble)')

        # ── Forecast start line ──
        if forecast_start_idx is not None and 0 < forecast_start_idx < n:
            ax.axvline(forecast_start_idx, color='#C62828', linewidth=1,
                       linestyle=':', alpha=0.6, zorder=2)
            ax.text(forecast_start_idx + 0.3, ax.get_ylim()[1] * 0.98,
                    'Forecast zone', fontsize=7, color='#C62828',
                    va='top', ha='left', fontstyle='italic')

        # ── X-axis dates ──
        n_ticks = min(10, n)
        tick_step = max(1, n // n_ticks)
        tick_pos = list(range(0, n, tick_step))
        if (n - 1) not in tick_pos:
            tick_pos.append(n - 1)
        ax.set_xticks(tick_pos)
        dlabels = []
        for i in tick_pos:
            d = dates[i]
            dlabels.append(d.strftime('%d/%m') if hasattr(d, 'strftime') else str(d)[-5:])
        ax.set_xticklabels(dlabels, fontsize=7, rotation=30, ha='right')

        ax.set_ylabel('Giá (VND)', fontsize=9)
        ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda v, _: f'{v:,.0f}'))
        ax.legend(loc='upper left', fontsize=7, framealpha=0.9, ncol=2)
        ax.grid(axis='both', alpha=0.3)
        ax.set_title(f'{symbol} — ML Forecast (XGBoost / LSTM / Ensemble)',
                     fontsize=11, fontweight='bold', color='#1F4E79', pad=10)

        # ── Metrics box: RMSE per model in forecast zone ──
        metrics_parts = []
        actual_arr = np.array(actual, dtype=float)
        for pred, name in [(xgb, 'XGB'), (lstm, 'LSTM'), (ens, 'Ens')]:
            if pred is None:
                continue
            mask = ~(np.isnan(actual_arr) | np.isnan(pred))
            if forecast_start_idx is not None:
                fs_mask = np.zeros(n, dtype=bool)
                fs_mask[forecast_start_idx:] = True
                mask = mask & fs_mask
            if mask.sum() > 0:
                rmse = np.sqrt(np.mean((actual_arr[mask] - pred[mask])**2))
                rmse_pct = rmse / np.nanmean(actual_arr[mask]) * 100
                metrics_parts.append(f'{name}: {rmse_pct:.1f}%')
        if metrics_parts:
            ax.text(0.98, 0.02, 'RMSE  ' + '  |  '.join(metrics_parts),
                    transform=ax.transAxes, fontsize=7.5, ha='right', va='bottom',
                    bbox=dict(boxstyle='round,pad=0.3', facecolor='white',
                              edgecolor='#CCCCCC', alpha=0.9))

        fig.tight_layout()
        return cls._to_buf(fig)

    @classmethod
    def top_vs_bottom(cls, top_syms, top_scores, bot_syms, bot_scores):
        """Side-by-side bar: Top 5 (green) vs Bottom 5 (red)."""
        plt.rcParams.update(CHART_STYLE)
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9, 3.2), gridspec_kw={'width_ratios': [1, 1]})

        # Top 5
        y = np.arange(len(top_syms))
        colors_top = ['#1B5E20' if s>=80 else '#2E7D32' if s>=65 else '#F57F17' if s>=50 else '#E65100' for s in top_scores]
        bars = ax1.barh(y, top_scores, color=colors_top, height=0.6, edgecolor='white')
        ax1.set_yticks(y); ax1.set_yticklabels(top_syms, fontsize=10, fontweight='bold')
        ax1.set_xlim(0, 105); ax1.invert_yaxis()
        for bar, sc in zip(bars, top_scores):
            ax1.text(bar.get_width()+1, bar.get_y()+bar.get_height()/2, str(int(sc)),
                     va='center', fontsize=10, fontweight='bold')
        ax1.set_title('TOP 5 — Tốt nhất', fontsize=11, fontweight='bold', color='#1B5E20')
        ax1.grid(axis='x', alpha=0.3)

        # Bottom 5
        y2 = np.arange(len(bot_syms))
        colors_bot = ['#B71C1C' if s<30 else '#C62828' if s<40 else '#E65100' if s<50 else '#F57F17' for s in bot_scores]
        bars2 = ax2.barh(y2, bot_scores, color=colors_bot, height=0.6, edgecolor='white')
        ax2.set_yticks(y2); ax2.set_yticklabels(bot_syms, fontsize=10, fontweight='bold')
        ax2.set_xlim(0, 105); ax2.invert_yaxis()
        for bar, sc in zip(bars2, bot_scores):
            ax2.text(bar.get_width()+1, bar.get_y()+bar.get_height()/2, str(int(sc)),
                     va='center', fontsize=10, fontweight='bold')
        ax2.set_title('BOTTOM 5 — Cần tránh', fontsize=11, fontweight='bold', color='#B71C1C')
        ax2.grid(axis='x', alpha=0.3)

        fig.suptitle('So sánh Top 5 vs Bottom 5', fontsize=13, fontweight='bold', color='#1F4E79', y=1.02)
        fig.tight_layout()
        return cls._to_buf(fig)

# ═══════════════════════════════════════════════════════════════
# EXCEL READER
# ═══════════════════════════════════════════════════════════════
COL_MAP = {
    'Symbol':'symbol','Score':'score','Rating':'rating',
    'Timing':'timing','HoldPlan':'hold_plan',
    'VNI':'vni_regime','Screener':'screener',
    'Sharpe':'sharpe','WinRate':'win_rate',
    'MaxDD':'max_dd','AnnRet':'ann_ret',
    'HMM':'hmm','VolReg':'vol_regime',
    'Forecast':'forecast','EnsRet%':'ens_ret_pct',
    'Conf':'confidence','MCVol':'mc_vol',
    'LockDD%':'lock_dd','PLoss3%':'prob_loss_3',
    'Entry':'entry','SL':'stop_loss',
    'TP2R':'tp_2r','TP3R':'tp_3r',
    'Shares':'shares','Value':'value',
    'Method':'method','Analysis':'analysis',
    'Sortino':'sortino','Calmar':'calmar',
    'VolRatio':'vol_ratio','ER':'er','CMF':'cmf',
    'RSI':'rsi','RS20d':'rs_20d','Slope':'slope',
    'Alpha':'alpha_ann','Beta':'beta',
    'VaR95':'var95','CVaR':'cvar',
    'GARCHPersist':'garch_persist','Kurtosis':'kurtosis',
    'Kelly':'kelly','HalfKelly':'half_kelly',
    'Edge':'edge','ProfitFactor':'profit_factor',
}

def _s(text):
    """Strip emoji from text."""
    for e in ['🟢','🔴','🟡','⚠️','✅','❌','⭐','📈','📉','🔥','💰','⬆️','⬇️','ℹ️']:
        text = str(text).replace(e,'')
    return text.strip()

def read_excel(path, top_n=5):
    df = pd.read_excel(path, sheet_name='Summary')
    rename = {}
    for orig, target in COL_MAP.items():
        for col in df.columns:
            if col == orig or col.lower() == orig.lower():
                rename[col] = target; break
    df = df.rename(columns=rename)
    num_cols = [v for v in COL_MAP.values() if v not in
                ('symbol','rating','timing','hold_plan','vni_regime','screener',
                 'hmm','vol_regime','forecast','mc_vol','method','analysis')]
    for col in num_cols:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)
    df_sorted = df.sort_values('score', ascending=False)
    df_top = df_sorted.head(top_n).reset_index(drop=True)
    df_bottom = df_sorted.tail(top_n).sort_values('score', ascending=True).reset_index(drop=True)
    n_total = len(df_sorted)

    # ── Load Forecast sheet (optional) ──
    # Expected columns: Symbol, Date, Close, Predicted, Upper (opt), Lower (opt)
    forecast_data = {}
    try:
        df_fc = pd.read_excel(path, sheet_name='Forecast')
        # Normalize column names
        fc_rename = {}
        for col in df_fc.columns:
            cl = col.strip().lower()
            if cl in ('symbol','ticker','mã'): fc_rename[col] = 'symbol'
            elif cl in ('date','ngày','time'): fc_rename[col] = 'date'
            elif cl in ('close','price','giá','actual'): fc_rename[col] = 'close'
            elif cl in ('predicted','forecast','dự báo','pred'): fc_rename[col] = 'predicted'
            elif cl in ('upper','upper_ci','upper_band'): fc_rename[col] = 'upper'
            elif cl in ('lower','lower_ci','lower_band'): fc_rename[col] = 'lower'
            elif cl in ('forecast_start','oos_start'): fc_rename[col] = 'forecast_start'
            elif cl in ('xgb','xgb_pred','xgboost','xgboost_pred'): fc_rename[col] = 'xgb_pred'
            elif cl in ('lstm','lstm_pred'): fc_rename[col] = 'lstm_pred'
            elif cl in ('ens','ens_pred','ensemble','ensemble_pred'): fc_rename[col] = 'ens_pred'
            elif cl in ('ci_upper','ens_upper'): fc_rename[col] = 'ci_upper'
            elif cl in ('ci_lower','ens_lower'): fc_rename[col] = 'ci_lower'
        df_fc = df_fc.rename(columns=fc_rename)
        if 'date' in df_fc.columns:
            df_fc['date'] = pd.to_datetime(df_fc['date'], errors='coerce')
        if 'symbol' in df_fc.columns and 'close' in df_fc.columns and 'predicted' in df_fc.columns:
            for sym, grp in df_fc.groupby('symbol'):
                grp = grp.sort_values('date').reset_index(drop=True)
                fc_entry = {
                    'dates': grp['date'].tolist(),
                    'close': grp['close'].tolist(),
                    'predicted': grp['predicted'].tolist(),
                }
                if 'upper' in grp.columns:
                    fc_entry['upper'] = grp['upper'].tolist()
                if 'lower' in grp.columns:
                    fc_entry['lower'] = grp['lower'].tolist()
                # ML model predictions
                for ml_col in ('xgb_pred', 'lstm_pred', 'ens_pred', 'ci_upper', 'ci_lower'):
                    if ml_col in grp.columns:
                        fc_entry[ml_col] = grp[ml_col].tolist()
                # Detect forecast_start: either from column or first row where close is NaN
                if 'forecast_start' in grp.columns:
                    fs_vals = grp['forecast_start'].dropna()
                    if len(fs_vals) > 0:
                        fc_entry['forecast_start_idx'] = int(fs_vals.iloc[0])
                else:
                    # Infer: last index where close is not NaN
                    close_arr = grp['close'].values
                    nan_mask = pd.isna(close_arr)
                    if nan_mask.any():
                        fc_entry['forecast_start_idx'] = int(np.argmax(nan_mask))
                forecast_data[str(sym).strip().upper()] = fc_entry
            log.info(f"Forecast data loaded for {len(forecast_data)} symbols: {list(forecast_data.keys())}")
        else:
            log.warning("Forecast sheet exists but missing required columns (symbol, close, predicted)")
    except Exception as e:
        log.info(f"No Forecast sheet found (optional): {e}")

    return df_top, df_bottom, n_total, forecast_data

# ═══════════════════════════════════════════════════════════════
# DOCX HELPERS
# ═══════════════════════════════════════════════════════════════
def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_paragraph(doc_or_cell, text, size=11, bold=False, color=CLR_DARK,
                          alignment=None, space_before=0, space_after=0, italic=False):
    """Add a paragraph with styling."""
    p = doc_or_cell.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = 'Arial'
    run.font.italic = italic
    # Force Arial for East Asian text too
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Arial"/>')
    rPr.append(rFonts)
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return p

def add_heading_blue(doc, text, level=1):
    """Add a heading with blue color."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = CLR_BLUE if level <= 2 else CLR_ACCENT
        run.font.name = 'Arial'
    return h

def make_table(doc, headers, rows_data, col_widths=None):
    """Create a formatted table."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows_data), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.name = 'Arial'
        run.font.color.rgb = CLR_DARK
        set_cell_shading(cell, 'D5E8F0')

    # Data rows
    for r_idx, row in enumerate(rows_data):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Arial'
            # Color assessment column
            val_str = str(val)
            if val_str.startswith('✅'):
                run.font.color.rgb = CLR_GREEN
            elif val_str.startswith('❌'):
                run.font.color.rgb = CLR_RED
            elif val_str.startswith('⚠'):
                run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
            else:
                run.font.color.rgb = CLR_DARK
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F5F5F5')

    # Set column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def add_alert_box(doc, text, border_color='0C5460', bg_color='D1ECF1', text_color=None):
    """Add an alert/callout box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = ''
    set_cell_shading(cell, bg_color)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor.from_string(text_color or border_color)
    # Set left border thick
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:color="{border_color}"/>'
        f'  <w:top w:val="single" w:sz="4" w:color="{bg_color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:color="{bg_color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:color="{bg_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    return table

# ═══════════════════════════════════════════════════════════════
# ANALYSIS ENGINE
# ═══════════════════════════════════════════════════════════════
def _hmm_vi(h):
    return _s(h).upper() or 'N/A'

def _vol_vi(v):
    v = _s(v).upper()
    return {'LOW':'Thấp','MEDIUM':'Trung bình','HIGH':'Cao',
            'EXTREME':'Cực cao','CALM':'Ổn định',
            'EXPANSION':'EXPANSION','NORMAL':'NORMAL',
            'CONTRACTION':'CONTRACTION'}.get(v, v)

def build_kpi_rows(r):
    rows = []
    ar=r.get('ann_ret',0); sh=r.get('sharpe',0); wr=r.get('win_rate',0)
    mdd=r.get('max_dd',0); so=r.get('sortino',0); ca=r.get('calmar',0)

    rows.append(['Hiệu suất','AnnRet',f'{ar:.2f}%',
        '✅ Xuất sắc' if ar>50 else '✅ Tốt' if ar>20 else '⚠️ Trung bình'])
    rows.append(['Hiệu suất','Sharpe',f'{sh:.3f}',
        '✅ Hàng đầu' if sh>2 else '✅ Tốt' if sh>=1 else '⚠️ Dưới 1.0'])
    rows.append(['Hiệu suất','WinRate',f'{wr:.1f}%',
        '✅ Tốt' if wr>=55 else '✅ Trên 50%' if wr>=50 else '⚠️ Thấp hơn 50%'])
    rows.append(['Hiệu suất','MaxDD',f'{mdd:.2f}%',
        '✅ Kiểm soát được' if mdd>-20 else '⚠️ Nặng' if mdd>-30 else '❌ Nghiêm trọng'])
    if so: rows.append(['Hiệu suất','Sortino',f'{so:.3f}',
        '✅ Tốt' if so>sh else '⚠️'])
    if ca: rows.append(['Hiệu suất','Calmar',f'{ca:.1f}',
        '✅ Tốt' if ca>=3 else '⚠️'])

    hmm = _hmm_vi(str(r.get('hmm','N/A'))); vreg = _vol_vi(str(r.get('vol_regime','N/A')))
    rows.append(['Thị trường','HMM',hmm,
        '✅ Xu hướng tăng rõ' if 'BULL' in hmm else '⚠️ Chưa rõ hướng' if 'SIDE' in hmm or 'NEUTRAL' in hmm else '❌ Giảm'])
    rows.append(['Thị trường','VolReg',vreg,
        '✅ Ổn định' if vreg in ('NORMAL','Thấp','Ổn định','CONTRACTION') else '⚠️ Biến động mở rộng'])

    er=r.get('ens_ret_pct',0); conf=r.get('confidence',0)
    rows.append(['Dự báo','Forecast','TĂNG' if er>0 else 'GIẢM', '✅' if er>0 else '❌'])
    rows.append(['Dự báo','EnsRet%',f'{er:+.2f}%', '✅' if er>1 else '⚠️' if er>0 else '❌'])
    rows.append(['Dự báo','Conf',f'{conf:.0f}%',
        '✅ Đồng thuận cao' if conf>=67 else '⚠️ Chưa đồng thuận cao' if conf>=50 else '❌ Thấp'])

    ldd=r.get('lock_dd',0); pl3=r.get('prob_loss_3',0)
    rows.append(['Rủi ro T+3','LockDD%',f'{ldd:.2f}%',
        '✅ An toàn' if ldd>-5 else '⚠️ Khá cao' if ldd>-10 else '❌ Rất cao'])
    rows.append(['Rủi ro T+3','PLoss3%',f'{pl3:.1f}%',
        '✅ Thấp' if pl3<20 else '⚠️ Đáng chú ý' if pl3<30 else '❌ Rủi ro lỗ nặng cao'])

    v95=r.get('var95',0); cv=r.get('cvar',0)
    if v95: rows.append(['VaR','VaR 95%',f'{v95:.2f}%/ngày',
        '✅ Thấp' if v95>-2 else '⚠️ Cao' if v95>-3 else '❌ Rất cao'])
    if cv: rows.append(['VaR','CVaR',f'{cv:.2f}%/ngày',
        '✅ Kiểm soát được' if cv>-3 else '⚠️ Tail risk đáng kể' if cv>-5 else '❌ Tail risk rất lớn'])
    return rows

def gen_commentary(r):
    """Generate prose-style commentary for a stock — approved template."""
    sym=r.get('symbol','?'); score=int(r.get('score',0))
    ar=r.get('ann_ret',0); sh=r.get('sharpe',0)
    wr=r.get('win_rate',0); mdd=r.get('max_dd',0); so=r.get('sortino',0)
    hmm=_hmm_vi(str(r.get('hmm','N/A')))
    vreg=_vol_vi(str(r.get('vol_regime','N/A'))); er=r.get('ens_ret_pct',0)
    conf=r.get('confidence',0); ldd=r.get('lock_dd',0); pl3=r.get('prob_loss_3',0)
    kur=r.get('kurtosis',0)
    entry=r.get('entry',0); sl=r.get('stop_loss',0)
    tp2=r.get('tp_2r',0); tp3=r.get('tp_3r',0)
    vni=_s(str(r.get('vni_regime','N/A'))).upper()
    rating=_s(str(r.get('rating','')))

    paras = []

    # ── Paragraph 1: Trạng thái (Score + HMM + VolReg) ──
    p1 = f'Score {score}, HMM (trạng thái xu hướng) đang {hmm} và VolReg (trạng thái biến động) ở {vreg}'
    if 'BULL' in hmm:
        p1 += ' — cổ phiếu đang trong pha tăng trưởng rõ ràng, lực đẩy mạnh'
    elif 'SIDE' in hmm or 'NEUTRAL' in hmm:
        p1 += ' — chưa xác lập rõ hướng đi, nhưng nền dao động'
        p1 += ' ổn định' if vreg in ('NORMAL','Ổn định','Thấp','CONTRACTION') else ' đang mở rộng'
    else:
        p1 += ' — xu hướng giảm được xác nhận bởi mô hình xác suất'
    if vreg == 'EXPANSION':
        p1 += ', tuy nhiên biên độ dao động đang mở rộng nên cần kiểm soát vị thế chặt hơn bình thường.'
    elif vreg in ('NORMAL','Ổn định','Thấp','CONTRACTION'):
        if 'BULL' in hmm:
            p1 += ', nền dao động ổn định — tổ hợp khá thuận lợi.'
        else:
            p1 += '.'
    else:
        p1 += '.'
    paras.append(p1)

    # ── Paragraph 2: Hiệu suất (Sharpe + Sortino + WinRate + AnnRet + MaxDD) ──
    # Sharpe assessment
    if sh >= 2:
        sh_assess = f'mức xuất sắc, cao nhất hạng: mỗi đơn vị rủi ro bạn chấp nhận, cổ phiếu trả lại hơn {sh:.1f} đơn vị lợi nhuận'
    elif sh >= 1.5:
        sh_assess = f'mức khá mạnh: mỗi đơn vị rủi ro bạn chấp nhận, cổ phiếu trả lại hơn {sh:.1f} đơn vị lợi nhuận'
    elif sh >= 1:
        sh_assess = 'mức tốt — lợi nhuận đang bù đắp xứng đáng cho rủi ro chấp nhận'
    elif sh >= 0.5:
        sh_assess = 'mức trung bình — lợi nhuận chưa thực sự nổi bật so với rủi ro'
    elif sh >= 0:
        sh_assess = 'mức thấp — lợi nhuận chưa đủ bù đắp rủi ro chấp nhận'
    else:
        sh_assess = 'mức âm — lợi nhuận không bù đắp được rủi ro, chiến lược đang lỗ trên cơ sở điều chỉnh rủi ro'

    p2 = f'Sharpe {sh:.2f} (lợi nhuận trên rủi ro) — trên 1.0 được coi là tốt, trên 2.0 là xuất sắc, nên {sh:.2f} nằm ở {sh_assess}.'

    # Sortino
    if so and so > sh:
        p2 += f' Sortino vượt Sharpe (chỉ tính rủi ro giảm) — nghĩa là phần lớn dao động đến từ chiều tăng giá, không phải chiều giảm, đây là dấu hiệu rất tích cực.'
    elif so and so > 0:
        p2 += f' Sortino {so:.2f} (chỉ tính rủi ro giảm) — chưa vượt Sharpe, biến động hai chiều khá cân bằng.'

    # WinRate
    if wr >= 55:
        p2 += f' WinRate {wr:.1f}% (tỷ lệ giao dịch thắng) — mức tốt, chiến lược thắng nhiều hơn thua rõ rệt.'
    elif wr >= 50:
        p2 += f' WinRate {wr:.1f}% (tỷ lệ giao dịch thắng) — trên ngưỡng 50% nghĩa là chiến lược thắng nhiều hơn thua'
        if sh >= 1:
            p2 += ', kết hợp với Sharpe cao cho thấy không chỉ thắng nhiều mà mỗi lần thắng còn lãi xứng đáng.'
        else:
            p2 += '.'
    else:
        p2 += f' WinRate {wr:.1f}% (tỷ lệ giao dịch thắng) — dưới 50%, chiến lược thắng ít trận nhưng cần mỗi trận thắng phải lãi lớn để bù lại.'

    # AnnRet
    if ar > 50:
        p2 += f' AnnRet {ar:.1f}% ấn tượng'
    elif ar > 20:
        p2 += f' AnnRet {ar:.1f}% chấp nhận được'
    elif ar > 0:
        p2 += f' AnnRet {ar:.1f}% chưa nổi bật'
    else:
        p2 += f' AnnRet {ar:.1f}% — lợi nhuận âm'

    # MaxDD
    if mdd > -15:
        p2 += f', MaxDD {mdd:.1f}% kiểm soát tốt.'
    elif mdd > -25:
        p2 += f', tuy nhiên MaxDD {mdd:.1f}% nhắc nhở: nếu vào sai đỉnh, mức thua tạm thời có thể tới 1/4 vốn trước khi hồi phục.'
    else:
        p2 += f', tuy nhiên MaxDD {mdd:.1f}% là mức nghiêm trọng — nếu vào sai timing, mức thua tạm thời có thể rất sâu.'
    paras.append(p2)

    # ── Paragraph 3: Dự báo (EnsRet + Conf + mô hình) ──
    p3 = f'EnsRet {er:+.1f}% (lợi nhuận kỳ vọng) cùng Conf {conf:.0f}% (độ đồng thuận mô hình) — pipeline chạy 3 mô hình độc lập (ARIMA, GARCH, Monte Carlo), mỗi mô hình nhìn giá từ một góc khác nhau.'
    if conf >= 67:
        p3 += f' Conf {conf:.0f}% nghĩa là 2/3 mô hình đồng ý hướng {"tăng" if er>0 else "giảm"} — ngưỡng từ 67% trở lên được coi là đáng tin cậy. Ở đây tín hiệu đang nghiêng khá rõ về phía {"tích cực" if er>0 else "tiêu cực"}.'
    elif conf >= 50:
        p3 += f' Conf {conf:.0f}% — các mô hình chưa thực sự đồng thuận, tín hiệu đang hình thành chứ chưa chín. Dưới 67% thì cần chờ thêm xác nhận trước khi hành động.'
    else:
        p3 += f' Conf {conf:.0f}% — thấp, các module cho tín hiệu trái chiều, không nên dựa vào dự báo này để ra quyết định.'
    paras.append(p3)

    # ── Paragraph 4: Rủi ro (LockDD + PLoss3 + VaR/CVaR/Kurtosis) ──
    p4 = f'Về rủi ro, LockDD {ldd:.2f}% (drawdown T+3) nghĩa là trong 3 ngày khóa lệnh, giá có thể giảm tới {abs(ldd):.1f}% trước khi bạn bán được'
    if ldd > -5:
        p4 += ' — mức an toàn, gần như yên tâm trong giai đoạn khóa lệnh.'
    elif ldd > -8:
        p4 += ' — mức khá cao, cần tính vào khi quyết định size.'
    else:
        p4 += ' — mức rất cao, đây là rủi ro đáng lo ngại.'

    p4 += f' PLoss3% {pl3:.1f}%'
    if pl3 < 20:
        p4 += ' — xác suất lỗ nặng thấp, khá an toàn.'
    elif pl3 < 30:
        p4 += f' — cứ 10 lần vào lệnh thì khoảng {pl3/10:.0f} lần sẽ lỗ vượt 3% trong giai đoạn khóa, đây là xác suất đáng chú ý.'
    else:
        p4 += f' — cứ 3 lần vào lệnh thì khoảng 1 lần sẽ lỗ vượt 3% trong giai đoạn khóa, đây là xác suất đáng cảnh giác.'

    # VaR/CVaR/Kurtosis
    v95 = r.get('var95', 0); cv = r.get('cvar', 0)
    tail_parts = []
    if v95: tail_parts.append(f'VaR (mức lỗ thông thường)')
    if cv: tail_parts.append(f'CVaR (lỗ trung bình tệ nhất)')
    if kur:
        tail_parts.append(f'Kurtosis (độ dày đuôi)')

    if tail_parts:
        p4 += ' ' + ', '.join(tail_parts)
        if kur > 4:
            p4 += f' cho thấy phân phối đuôi khá dày (Kurtosis {kur:.1f} > 4) — rủi ro cực đoan cao hơn bình thường, cần đề phòng sự kiện "thiên nga đen".'
        elif kur > 3:
            p4 += ' cho thấy đuôi hơi dày nhưng chưa đến mức đáng ngại — rủi ro cực đoan ở mức kiểm soát được.'
        else:
            p4 += ' chưa cho thấy rủi ro cực đoan — phân phối lợi nhuận chưa có đuôi dày bất thường.'
    paras.append(p4)

    # ── Paragraph 5: Kết luận hành động ──
    # Build conclusion
    p5 = 'Kết luận hành động: '
    if score >= 65 and sh >= 1 and conf >= 67:
        p5 += f'{sym} dẫn đầu bảng xếp hạng với nền tảng quant mạnh và tín hiệu đồng thuận cao.'
    elif score >= 60 and sh >= 1:
        p5 += f'{sym} có nền tảng quant tích cực, hiệu suất điều chỉnh rủi ro tốt.'
    elif score >= 50:
        p5 += f'{sym} có tín hiệu hỗn hợp — một số chỉ số tích cực nhưng chưa đủ đồng bộ.'
    else:
        p5 += f'{sym} có điểm quant thấp, đa số chỉ số chưa thuyết phục.'

    # Add caveats
    caveats = []
    if 'BEAR' in vni:
        caveats.append('VNI đang BEAR')
    if pl3 > 30:
        caveats.append('rủi ro T+3 khá cao')
    if 'SIDE' in hmm or 'NEUTRAL' in hmm:
        caveats.append(f'HMM {hmm} chưa xác nhận xu hướng')
    if mdd < -30:
        caveats.append(f'MaxDD {mdd:.1f}% nghiêm trọng')
    if caveats:
        p5 += ' Tuy nhiên ' + ' và '.join(caveats) + ', nên '
    else:
        p5 += ' '

    # Action
    if entry > 0 and sl > 0 and sl != entry:
        rpct = abs(entry - sl) / entry * 100
        rr2 = (tp2 - entry) / abs(entry - sl) if tp2 > 0 and entry != sl else 0
        rr3 = (tp3 - entry) / abs(entry - sl) if tp3 > 0 and entry != sl else 0

        if pl3 > 35 or 'KHÔNG' in _s(str(r.get('timing',''))).upper():
            p5 += f'không nên vào lệnh ở thời điểm hiện tại. Chờ tín hiệu cải thiện trước khi hành động.'
        elif 'PULLBACK' in _s(str(r.get('timing',''))).upper() or pl3 > 25:
            p5 += f'không nên lao vào mua ngay. Chờ giá pullback về vùng entry {entry:,.0f}'
            p5 += f', đặt cắt lỗ kỷ luật tại {sl:,.0f} (rủi ro -{rpct:.1f}%)'
            if tp2 > 0:
                p5 += f', target chốt lời tại {tp2:,.0f} (R:R 1:{rr2:.1f})'
            if tp3 > 0:
                p5 += f' hoặc {tp3:,.0f} (R:R 1:{rr3:.1f})'
            p5 += '.'
        else:
            p5 += f'Entry tại {entry:,.0f}, cắt lỗ kỷ luật tại {sl:,.0f} (rủi ro -{rpct:.1f}%)'
            if tp2 > 0:
                p5 += f', target {tp2:,.0f} (R:R 1:{rr2:.1f})'
            if tp3 > 0:
                p5 += f' hoặc {tp3:,.0f} (R:R 1:{rr3:.1f})'
            p5 += '.'
    else:
        p5 += 'Chưa xác định được entry/SL phù hợp — tiếp tục theo dõi.'
    paras.append(p5)

    return paras


def gen_bottom_commentary(r):
    """Generate prose-style commentary for bottom stocks — why to avoid."""
    sym=r.get('symbol','?'); score=int(r.get('score',0))
    ar=r.get('ann_ret',0); sh=r.get('sharpe',0)
    wr=r.get('win_rate',0); mdd=r.get('max_dd',0)
    hmm=_hmm_vi(str(r.get('hmm','N/A')))
    vreg=_vol_vi(str(r.get('vol_regime','N/A'))); er=r.get('ens_ret_pct',0)
    conf=r.get('confidence',0); ldd=r.get('lock_dd',0); pl3=r.get('prob_loss_3',0)
    kur=r.get('kurtosis',0)

    paras = []

    # ── P1: Trạng thái ──
    p1 = f'Score {score}, HMM (trạng thái xu hướng) đang {hmm} và VolReg (trạng thái biến động) ở {vreg}'
    if 'BEAR' in hmm:
        p1 += ' — mô hình xác suất xác nhận xu hướng giảm rõ ràng, đây là tín hiệu tiêu cực mạnh.'
    elif 'SIDE' in hmm or 'NEUTRAL' in hmm:
        p1 += ' — chưa có xu hướng rõ, nhưng với score thấp thì không đáng mạo hiểm.'
    else:
        p1 += ' — xu hướng tăng nhưng các chỉ số khác quá yếu để tin tưởng.'
    paras.append(p1)

    # ── P2: Hiệu suất ──
    p2 = f'Sharpe {sh:.2f} (lợi nhuận trên rủi ro)'
    if sh < 0:
        p2 += f' — mức âm, nghĩa là chiến lược đang lỗ trên cơ sở điều chỉnh rủi ro. Mỗi đơn vị rủi ro bạn chấp nhận, cổ phiếu lấy đi thay vì trả lại.'
    elif sh < 0.5:
        p2 += ' — mức quá thấp, lợi nhuận không bù đắp nổi rủi ro.'
    else:
        p2 += ' — mức trung bình nhưng không đủ để bù đắp cho các yếu tố tiêu cực khác.'
    p2 += f' WinRate {wr:.1f}% (tỷ lệ giao dịch thắng)'
    if wr < 45:
        p2 += ' — quá thấp, đa số giao dịch đều thua lỗ.'
    else:
        p2 += ' — chưa đủ thuyết phục.'
    if mdd < -30:
        p2 += f' MaxDD {mdd:.1f}% — mức sụt giảm nghiêm trọng, rủi ro "cháy tài khoản" rất cao.'
    elif mdd < -20:
        p2 += f' MaxDD {mdd:.1f}% — mức sụt giảm nặng, khó chịu đựng tâm lý.'
    paras.append(p2)

    # ── P3: Dự báo ──
    p3 = f'EnsRet {er:+.2f}% (lợi nhuận kỳ vọng) cùng Conf {conf:.0f}% (độ đồng thuận mô hình)'
    if er < 0:
        p3 += ' — mô hình dự báo giảm giá, không có lý do để kỳ vọng đảo chiều trong ngắn hạn.'
    elif er < 1:
        p3 += ' — gần như không có tín hiệu tăng rõ ràng, kỳ vọng lợi nhuận quá mỏng.'
    else:
        p3 += ' — có tín hiệu tăng nhẹ nhưng không đủ mạnh để bù đắp cho các rủi ro đã nêu.'
    if conf < 50:
        p3 += ' Các module cho tín hiệu trái chiều, không đáng tin cậy.'
    paras.append(p3)

    # ── P4: Rủi ro ──
    p4 = f'Về rủi ro, LockDD {ldd:.2f}% (drawdown T+3) và PLoss3% {pl3:.1f}%'
    if pl3 > 30:
        p4 += f' — xác suất lỗ nặng trong giai đoạn khóa lệnh rất cao, gần {pl3:.0f}% khả năng lỗ vượt 3%.'
    elif pl3 > 20:
        p4 += ' — rủi ro T+3 ở mức đáng lo ngại.'
    else:
        p4 += ' — rủi ro T+3 chưa quá cao nhưng không đủ để cứu vãn các chỉ số khác.'
    paras.append(p4)

    # ── P5: Kết luận ──
    p5 = f'Kết luận: {sym} (Score {score}) — KHÔNG MUA. Đa số chỉ số đều nằm trong vùng nguy hiểm. Nếu đang nắm giữ, cân nhắc cắt lỗ tại mức hiện tại.'
    paras.append(p5)

    return paras

def gen_conclusion(r):
    sym=r.get('symbol','?'); score=r.get('score',0); sh=r.get('sharpe',0)
    pl3=r.get('prob_loss_3',0); hmm=_hmm_vi(str(r.get('hmm','')))
    timing=_s(str(r.get('timing',''))); er=r.get('ens_ret_pct',0); conf=r.get('confidence',0)
    entry=r.get('entry',0); sl=r.get('stop_loss',0); tp2=r.get('tp_2r',0); tp3=r.get('tp_3r',0)
    shares=r.get('shares',0); value=r.get('value',0); hold=_s(str(r.get('hold_plan','')))
    rsi=r.get('rsi',0); mdd=r.get('max_dd',0); vni=_s(str(r.get('vni_regime',''))).upper()

    # Conclusion text
    parts = [f'Kết luận: {sym} ']
    if sh>=2: parts.append('có nền tảng định lượng tốt nhất. ')
    elif sh>=1: parts.append('có nền tảng định lượng tích cực. ')
    else: parts.append('có hiệu suất trung bình. ')
    if pl3<20: parts.append('Rủi ro T+3 thấp. ')
    elif pl3<30: parts.append('Rủi ro T+3 trung bình. ')
    else: parts.append('Rủi ro T+3 cao — cần quản lý thận trọng. ')
    if 'SIDE' in hmm or 'NEUTRAL' in hmm: parts.append(f'HMM {hmm} đòi hỏi thêm xác nhận. ')
    if rsi>70: parts.append(f'RSI {rsi:.0f} quá mua — chờ pullback. ')
    if mdd<-30: parts.append(f'MaxDD {mdd:.1f}% đòi hỏi SL nghiêm ngặt. ')
    if 'BEAR' in vni: parts.append('Trong bối cảnh VNI BEAR, ưu tiên kỷ luật SL và quản lý size.')
    conclusion = ''.join(parts)

    # Priority
    if 'KHÔNG' in timing.upper() or pl3>40: action='THẤP — chờ thêm xác nhận'
    elif 'PULLBACK' in timing.upper() or rsi>70: action='TRUNG BÌNH — chờ pullback'
    elif 'THEO DÕI' in timing.upper(): action='TRUNG BÌNH — chờ xác nhận'
    else: action='CAO — theo kế hoạch' if score>=70 and sh>=1.5 else 'TRUNG BÌNH'

    # Target
    rpct=abs(entry-sl)/entry*100 if entry>0 and sl>0 else 0
    rr2=(tp2-entry)/abs(entry-sl) if entry>0 and sl>0 and tp2>0 and entry!=sl else 0
    up2=(tp2-entry)/entry*100 if entry>0 and tp2>0 else 0
    up3=(tp3-entry)/entry*100 if entry>0 and tp3>0 else 0
    rr3=(tp3-entry)/abs(entry-sl) if entry>0 and sl>0 and tp3>0 and entry!=sl else 0

    target = [
        ['Entry', f'{entry:,.0f} VND' if entry>0 else 'N/A'],
        ['Stop Loss', f'{sl:,.0f} VND (-{rpct:.1f}%)' if sl>0 else 'Cần xác định thêm'],
        ['Target 1 (TP2R)', f'{tp2:,.0f} VND (+{up2:.1f}%, R:R 1:{rr2:.1f})' if tp2>0 else 'N/A'],
        ['Target 2 (TP3R)', f'{tp3:,.0f} VND (+{up3:.1f}%, R:R 1:{rr3:.1f})' if tp3>0 else 'N/A'],
        ['Kích thước', f'{shares:,.0f} cổ phần (~{value/1e6:.0f} triệu VND)' if shares>0 else 'Chưa xác định'],
        ['Thời gian nắm giữ', hold or '7 phiên (~10 ngày)'],
        ['Ưu tiên', action],
    ]
    return conclusion, target

# ═══════════════════════════════════════════════════════════════
# GLOSSARY DATA
# ═══════════════════════════════════════════════════════════════
GLOSSARY = [
    ('1.1 Chỉ số hiệu suất (Performance Metrics)', [
        ['Score','Điểm tổng hợp do pipeline tính trên thang 0–100','>= 60'],
        ['AnnRet (%)','Lợi nhuận hàng năm chuẩn hóa','> 20%'],
        ['Sharpe Ratio','Lợi nhuận thặng dư / Độ lệch chuẩn. Sharpe = (AnnRet - Rf) / Vol','>= 1.0 tốt; >= 2.0 xuất sắc'],
        ['Sortino','Như Sharpe nhưng chỉ tính biến động đi xuống','Sortino > Sharpe = tốt'],
        ['Calmar','Lợi nhuận năm / MaxDD','>= 3.0 = tốt'],
        ['WinRate (%)','Tỷ lệ phiên thắng trong backtest','>= 50%'],
        ['MaxDD (%)','Mức thua lỗ lớn nhất từ đỉnh xuống đáy','< -20% là nặng'],
    ]),
    ('1.2 Chỉ số thị trường (Market Statistics)', [
        ['HMM','Mô hình Markov ẩn — BULL / SIDEWAY / BEAR','BULL = vào lệnh'],
        ['VolReg','Chế độ biến động: EXPANSION / NORMAL / CONTRACTION','EXPANSION = cẩn thận'],
        ['CMF','Dòng tiền Chaikin — dương = tích lũy, âm = phân phối','> 0 = tích lũy'],
        ['RSI','RSI > 70 = quá mua, < 30 = quá bán','30–70 = trung tính'],
    ]),
    ('1.3 Chỉ số dự báo (Forecasting)', [
        ['EnsRet%','Lợi nhuận kỳ vọng từ ensemble nhiều mô hình','Càng cao càng tốt'],
        ['Conf (%)','Mức đồng thuận giữa các mô hình','>= 67% = khá tin cậy'],
        ['MCVol','Monte Carlo biến động: GARCH hoặc Historical','GARCH chính xác hơn'],
    ]),
    ('1.4 Rủi ro T+3 (Lock Risk)', [
        ['LockDD%','Drawdown tối đa kỳ vọng trong T+3 lock','Càng gần 0 càng tốt'],
        ['PLoss3%','Xác suất lỗ vượt 3% trong T+3','< 20% = an toàn'],
    ]),
    ('1.5 VaR / CVaR', [
        ['VaR 95%','Mức lỗ tối đa 95% các phiên','Biết mức lỗ "thông thường"'],
        ['CVaR','Trung bình mức lỗ 5 ngày tệ nhất / 100 ngày','CVaR luôn tệ hơn VaR'],
        ['Kurtosis','Độ nhọn phân phối. > 3 = đuôi dày','> 4 = cẩn thận "thiên nga đen"'],
    ]),
    ('1.6 Giao dịch (Trading Parameters)', [
        ['Entry / SL','Giá vào lệnh / Ngưỡng cắt lỗ','Bắt buộc kỷ luật'],
        ['TP2R / TP3R','Mục tiêu chốt lời tại R:R 1:2 và 1:3','Chốt lời theo kế hoạch'],
        ['Kelly','Kích thước vị thế tối ưu. Kelly âm = không vào lệnh','Kelly âm = không vào'],
    ]),
]

# ═══════════════════════════════════════════════════════════════
# MAIN DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════
def generate_report(excel_path='quant_report.xlsx', output_path='bao_cao_quant.docx', top_n=5):
    log.info(f"Reading {excel_path}")
    df, df_bot, n_total, forecast_data = read_excel(excel_path, top_n)
    log.info(f"Top {top_n}: {df['symbol'].tolist()}")
    log.info(f"Bottom {top_n}: {df_bot['symbol'].tolist()}")

    vni = _s(str(df['vni_regime'].iloc[0])) if 'vni_regime' in df.columns else 'N/A'
    date_str = datetime.now().strftime('%d/%m/%Y')

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ════════════════════════════════════════════════════
    # TITLE PAGE
    # ════════════════════════════════════════════════════
    add_styled_paragraph(doc, 'BÁO CÁO PHÂN TÍCH ĐỊNH LƯỢNG — TOP 5 CỔ PHIẾU',
                          size=18, bold=True, color=CLR_BLUE,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=40, space_after=12)
    add_styled_paragraph(doc, f'Hệ thống: quant_pipeline.py | Ngày: {date_str} | Bối cảnh VNI: {vni}',
                          size=11, color=CLR_GRAY,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    # ════════════════════════════════════════════════════
    # PART I: GLOSSARY
    # ════════════════════════════════════════════════════
    add_heading_blue(doc, 'PHẦN I — CHÚ THÍCH CÁC CHỈ SỐ', level=1)
    add_styled_paragraph(doc,
        'Trước khi đi vào phân tích từng mã, phần này giải thích ý nghĩa của tất cả các chỉ số '
        'định lượng sử dụng trong pipeline. Hiểu rõ từng chỉ số giúp nhà đầu tư đọc và diễn giải '
        'báo cáo một cách chính xác.',
        size=10, color=CLR_GRAY, space_after=8)

    for section_title, items in GLOSSARY:
        add_heading_blue(doc, section_title, level=3)
        make_table(doc, ['Chỉ số', 'Ý nghĩa', 'Ngưỡng tốt'], items, col_widths=[3, 9, 4])
        doc.add_paragraph()

    # ════════════════════════════════════════════════════
    # PART II: OVERVIEW
    # ════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading_blue(doc, 'PHẦN II — TỔNG QUAN TOP 5 CỔ PHIẾU', level=1)
    add_styled_paragraph(doc,
        f'Pipeline đã lọc và xếp hạng {n_total} mã cổ phiếu trong bối cảnh VNI {vni}. '
        f'Bảng sau tổng hợp 5 mã đứng đầu theo điểm tổng hợp:',
        size=10, color=CLR_GRAY, space_after=8)

    summary_rows = []
    for i, row in df.iterrows():
        er = row.get('ens_ret_pct',0)
        summary_rows.append([
            str(i+1), str(row.get('symbol','')), str(int(row.get('score',0))),
            _s(str(row.get('rating','')))[:25], _s(str(row.get('timing','')))[:18],
            f'{er:+.1f}%', f'{row.get("confidence",0):.0f}%',
            f'{row.get("entry",0):,.0f}', f'{row.get("stop_loss",0):,.0f}',
            f'{row.get("tp_2r",0):,.0f}', f'{row.get("tp_3r",0):,.0f}',
        ])
    make_table(doc,
        ['#','Mã','Score','Rating','Timing','Forecast','Conf','Entry','SL','TP2R','TP3R'],
        summary_rows, col_widths=[0.8, 1.2, 1, 3, 2.2, 1.5, 1.2, 1.8, 1.5, 1.5, 1.5])
    doc.add_paragraph()

    if 'BEAR' in vni.upper() or 'CRISIS' in vni.upper():
        add_alert_box(doc,
            f'Lưu ý quan trọng: Tất cả 5 mã đều thuộc bối cảnh VNI {vni}. '
            'Dù điểm cổ phiếu cao, rủi ro thị trường chung vẫn đang bất lợi. '
            'Ưu tiên kỷ luật stop loss và quản lý kích thước vị thế.',
            border_color='856404', bg_color='FFF3CD', text_color='856404')
        doc.add_paragraph()

    # ── INFOGRAPHIC: Score Ranking ──
    syms = df['symbol'].tolist(); scores = df['score'].tolist()
    doc.add_picture(Charts.score_ranking(syms, scores), width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── INFOGRAPHIC: Risk/Reward Map ──
    doc.add_picture(Charts.risk_reward_map(
        syms, df['entry'].tolist(), df['stop_loss'].tolist(),
        df['tp_2r'].tolist(), df['tp_3r'].tolist()), width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── INFOGRAPHIC: Performance Comparison ──
    doc.add_picture(Charts.performance_comparison(
        syms, df['sharpe'].tolist(), df['win_rate'].tolist(), df['max_dd'].tolist()), width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── INFOGRAPHIC: Risk Heatmap ──
    doc.add_picture(Charts.risk_heatmap(
        syms, df['prob_loss_3'].tolist(), df['lock_dd'].tolist(), df['confidence'].tolist()), width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ════════════════════════════════════════════════════
    # PART III: DETAILED ANALYSIS
    # ════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading_blue(doc, 'PHẦN III — PHÂN TÍCH CHI TIẾT TỪNG MÃ', level=1)

    for idx, row in df.iterrows():
        r = row.to_dict()
        for k,v in r.items():
            if isinstance(v, str): r[k] = _s(v)
            if isinstance(v, float) and np.isnan(v): r[k] = 0

        sym = r.get('symbol','?'); score = int(r.get('score',0))
        rating = r.get('rating',''); timing = r.get('timing','')

        if idx > 0:
            doc.add_page_break()

        add_heading_blue(doc, f'{idx+1}. {sym} — Điểm {score} | {rating}', level=2)
        add_styled_paragraph(doc, f'Timing: {timing}', size=10, color=CLR_GRAY, italic=True, space_after=6)

        # KPI table
        add_heading_blue(doc, 'Bảng chỉ số tổng hợp', level=3)
        kpi_rows = build_kpi_rows(r)
        make_table(doc, ['Nhóm','Chỉ số','Giá trị','Đánh giá'], kpi_rows, col_widths=[3, 2.5, 4, 6.5])
        doc.add_paragraph()

        # ── INFOGRAPHIC: Radar chart per stock ──
        sharpe = r.get('sharpe',0); wr = r.get('win_rate',0)
        er_val = r.get('ens_ret_pct',0); conf_val = r.get('confidence',0)
        pl3_val = r.get('prob_loss_3',0)
        factors = {
            'Sharpe': np.clip(sharpe / 2, 0, 1),
            'WinRate': np.clip((wr - 30) / 40, 0, 1),
            'Forecast': np.clip((er_val + 2) / 8, 0, 1),
            'Confidence': np.clip(conf_val / 100, 0, 1),
            'An toàn T+3': np.clip((50 - pl3_val) / 50, 0, 1),
        }
        radar_buf = Charts.stock_radar(sym, factors)
        if radar_buf:
            doc.add_picture(radar_buf, width=Cm(8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # ── INFOGRAPHIC: Price Forecast chart (if data available) ──
        sym_upper = str(sym).strip().upper()
        if sym_upper in forecast_data:
            fc = forecast_data[sym_upper]
            try:
                fc_buf = Charts.price_forecast(
                    symbol=sym,
                    dates=fc['dates'],
                    actual=fc['close'],
                    predicted=fc['predicted'],
                    upper=fc.get('upper'),
                    lower=fc.get('lower'),
                    forecast_start_idx=fc.get('forecast_start_idx'),
                )
                add_heading_blue(doc, 'Biểu đồ dự báo giá', level=3)
                doc.add_picture(fc_buf, width=Cm(16))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Add forecast metrics caption
                fc_close = np.array(fc['close'], dtype=float)
                fc_pred = np.array(fc['predicted'], dtype=float)
                mask = ~(np.isnan(fc_close) | np.isnan(fc_pred))
                if mask.sum() > 1:
                    mape = np.mean(np.abs((fc_close[mask] - fc_pred[mask]) / fc_close[mask])) * 100
                    a_dir = np.diff(fc_close[mask])
                    p_dir = np.diff(fc_pred[mask])
                    da = np.mean(np.sign(a_dir) == np.sign(p_dir)) * 100
                    last_pred = fc_pred[~np.isnan(fc_pred)][-1] if any(~np.isnan(fc_pred)) else 0
                    last_close = fc_close[~np.isnan(fc_close)][-1] if any(~np.isnan(fc_close)) else 0
                    pred_ret = (last_pred / last_close - 1) * 100 if last_close > 0 else 0
                    caption = (
                        f'MAPE (sai số trung bình): {mape:.1f}%  |  '
                        f'DA (chính xác hướng): {da:.0f}%  |  '
                        f'Dự báo tiếp theo: {last_pred:,.0f} ({pred_ret:+.1f}%)'
                    )
                    add_styled_paragraph(doc, caption, size=9, color=CLR_GRAY,
                                         italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
                doc.add_paragraph()
            except Exception as e:
                log.warning(f"Failed to generate forecast chart for {sym}: {e}")

            # ── ML Forecast chart (if multi-model data available) ──
            has_ml = any(k in fc for k in ('xgb_pred', 'lstm_pred', 'ens_pred'))
            if has_ml:
                try:
                    ml_buf = Charts.ml_forecast(
                        symbol=sym,
                        dates=fc['dates'],
                        actual=fc['close'],
                        xgb_pred=fc.get('xgb_pred'),
                        lstm_pred=fc.get('lstm_pred'),
                        ens_pred=fc.get('ens_pred'),
                        ci_upper=fc.get('ci_upper'),
                        ci_lower=fc.get('ci_lower'),
                        forecast_start_idx=fc.get('forecast_start_idx'),
                    )
                    add_heading_blue(doc, 'ML Forecast — XGBoost / LSTM / Ensemble', level=3)
                    doc.add_picture(ml_buf, width=Cm(16))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except Exception as e:
                    log.warning(f"Failed to generate ML forecast chart for {sym}: {e}")

        # Commentary (prose style)
        add_heading_blue(doc, 'Phân tích', level=3)
        commentary_paras = gen_commentary(r)
        for i, para_text in enumerate(commentary_paras):
            is_last = (i == len(commentary_paras) - 1)
            if is_last:
                # Last paragraph = conclusion action → alert box
                add_alert_box(doc, para_text, border_color='0C5460', bg_color='D1ECF1', text_color='0C5460')
            else:
                add_styled_paragraph(doc, para_text, size=10, color=CLR_DARK, space_after=6)

    # ════════════════════════════════════════════════════
    # PART IV: CONCLUSION
    # ════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading_blue(doc, 'PHẦN IV — KẾT LUẬN VÀ XẾP HẠNG ƯU TIÊN', level=1)

    # Summary ranking table
    add_heading_blue(doc, 'Tổng kết so sánh Top 5', level=3)
    rank_rows = []
    for idx, row in df.iterrows():
        r = row.to_dict()
        for k,v in r.items():
            if isinstance(v, str): r[k] = _s(v)
            if isinstance(v, float) and np.isnan(v): r[k] = 0
        _, target = gen_conclusion(r)
        priority = target[-1][1]  # last row = Ưu tiên
        rank_rows.append([
            r.get('symbol',''), str(int(r.get('score',0))),
            _s(str(r.get('timing','')))[:15], f'{r.get("sharpe",0):.3f}',
            f'{r.get("max_dd",0):.2f}%', f'{r.get("prob_loss_3",0):.1f}%',
            _hmm_vi(str(r.get('hmm',''))), f'{r.get("confidence",0):.0f}%', priority
        ])
    make_table(doc, ['Mã','Score','Timing','Sharpe','MaxDD','PLoss3%','HMM','Conf','Hành động'],
               rank_rows, col_widths=[1.5, 1, 2, 1.5, 1.8, 1.5, 1.5, 1, 4.2])
    doc.add_paragraph()

    if 'BEAR' in vni.upper() or 'CRISIS' in vni.upper():
        add_alert_box(doc,
            f'Bối cảnh quan trọng: Toàn bộ top 5 đều trong môi trường VNI {vni}. '
            'Chiến lược đúng đắn nhất hiện tại là phòng thủ và kiên nhẫn chờ điểm vào lệnh tối ưu.',
            border_color='856404', bg_color='FFF3CD', text_color='856404')
        doc.add_paragraph()

    # Rankings
    add_heading_blue(doc, 'Xếp hạng ưu tiên đầu tư', level=3)
    for idx, row in df.iterrows():
        r = row.to_dict()
        for k,v in r.items():
            if isinstance(v, str): r[k] = _s(v)
            if isinstance(v, float) and np.isnan(v): r[k] = 0
        conclusion, _ = gen_conclusion(r)
        add_styled_paragraph(doc, f'{idx+1}. {r.get("symbol","")} (Ưu tiên #{idx+1}): {conclusion}',
                              size=10, space_after=4)
    doc.add_paragraph()

    # Recommendations
    add_heading_blue(doc, 'Khuyến nghị quản lý danh mục', level=3)
    recs = [
        f'Phân bổ vốn: Trong {vni} market, không nên phân bổ quá 15–20% tổng danh mục vào mỗi mã.',
        'Stop loss bắt buộc: Kỷ luật cắt lỗ là yêu cầu tuyệt đối.',
        'Theo dõi VNI: Khi VNI chuyển sang NEUTRAL hoặc BULL, tín hiệu sẽ được nâng cấp đáng kể.',
        'Tái đánh giá: Pipeline nên chạy lại mỗi 3–5 phiên để cập nhật tín hiệu.',
    ]
    for r in recs:
        add_styled_paragraph(doc, r, size=10, color=CLR_GRAY, space_after=3)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════
    # PART V: BOTTOM 5 — CỔ PHIẾU CẦN TRÁNH
    # ════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading_blue(doc, 'PHẦN V — BOTTOM 5: CỔ PHIẾU CẦN TRÁNH', level=1)
    add_styled_paragraph(doc,
        f'Ngoài việc tìm kiếm cơ hội, việc nhận diện cổ phiếu yếu nhất cũng quan trọng không kém. '
        f'Dưới đây là 5 mã có điểm thấp nhất trong {n_total} mã được pipeline quét. '
        f'Đây là các mã mà nhà đầu tư KHÔNG nên vào lệnh ở thời điểm hiện tại.',
        size=10, color=CLR_GRAY, space_after=8)

    # ── Top vs Bottom comparison chart ──
    doc.add_picture(Charts.top_vs_bottom(
        df['symbol'].tolist(), df['score'].tolist(),
        df_bot['symbol'].tolist(), df_bot['score'].tolist()), width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── Bottom 5 summary table ──
    add_heading_blue(doc, 'Bảng tổng hợp Bottom 5', level=3)
    bot_rows = []
    for i, row in df_bot.iterrows():
        er = row.get('ens_ret_pct', 0)
        bot_rows.append([
            str(i+1), str(row.get('symbol','')), str(int(row.get('score',0))),
            _s(str(row.get('rating','')))[:25],
            f'{er:+.1f}%', f'{row.get("confidence",0):.0f}%',
            f'{row.get("sharpe",0):.2f}', f'{row.get("max_dd",0):.1f}%',
            f'{row.get("prob_loss_3",0):.1f}%',
            _hmm_vi(str(row.get('hmm','N/A'))),
        ])
    make_table(doc,
        ['#','Mã','Score','Rating','Forecast','Conf','Sharpe','MaxDD','PLoss3%','HMM'],
        bot_rows, col_widths=[0.8, 1.2, 1, 3, 1.5, 1.2, 1.5, 1.5, 1.5, 2.3])
    doc.add_paragraph()

    add_alert_box(doc,
        'Cảnh báo: Các mã dưới đây có điểm tổng hợp thấp nhất — đa số chỉ số đều ở vùng tiêu cực. '
        'Tuyệt đối không vào lệnh mua. Nếu đang nắm giữ, cân nhắc cắt lỗ hoặc giảm vị thế.',
        border_color='721C24', bg_color='F8D7DA', text_color='721C24')
    doc.add_paragraph()

    # ── Detailed analysis for each bottom stock ──
    for idx, row in df_bot.iterrows():
        r = row.to_dict()
        for k, v in r.items():
            if isinstance(v, str): r[k] = _s(v)
            if isinstance(v, float) and np.isnan(v): r[k] = 0

        sym = r.get('symbol', '?'); score = int(r.get('score', 0))
        rating = r.get('rating', '')
        sh = r.get('sharpe', 0); wr = r.get('win_rate', 0)
        mdd = r.get('max_dd', 0); ar = r.get('ann_ret', 0)
        er = r.get('ens_ret_pct', 0); conf = r.get('confidence', 0)
        pl3 = r.get('prob_loss_3', 0); ldd = r.get('lock_dd', 0)
        hmm = _hmm_vi(str(r.get('hmm', 'N/A')))
        vreg = _vol_vi(str(r.get('vol_regime', 'N/A')))

        add_heading_blue(doc, f'Bottom #{idx+1}: {sym} — Điểm {score} | {rating}', level=2)

        # KPI table
        kpi_rows = build_kpi_rows(r)
        make_table(doc, ['Nhóm','Chỉ số','Giá trị','Đánh giá'], kpi_rows, col_widths=[3, 2.5, 4, 6.5])
        doc.add_paragraph()

        # Radar chart
        factors = {
            'Sharpe': np.clip(sh / 2, 0, 1),
            'WinRate': np.clip((wr - 30) / 40, 0, 1),
            'Forecast': np.clip((er + 2) / 8, 0, 1),
            'Confidence': np.clip(conf / 100, 0, 1),
            'An toàn T+3': np.clip((50 - pl3) / 50, 0, 1),
        }
        radar_buf = Charts.stock_radar(sym, factors)
        if radar_buf:
            doc.add_picture(radar_buf, width=Cm(8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # Commentary (prose style)
        add_heading_blue(doc, 'Phân tích', level=3)
        bot_commentary = gen_bottom_commentary(r)
        for i, para_text in enumerate(bot_commentary):
            is_last = (i == len(bot_commentary) - 1)
            if is_last:
                add_alert_box(doc, para_text, border_color='721C24', bg_color='F8D7DA', text_color='721C24')
            else:
                add_styled_paragraph(doc, para_text, size=10, color=CLR_DARK, space_after=6)

        if idx < len(df_bot) - 1:
            doc.add_page_break()

    # Disclaimer
    add_styled_paragraph(doc, f'Báo cáo được tạo tự động từ quant_pipeline.py — Ngày {date_str}',
                          size=9, color=CLR_GRAY, italic=True)
    add_styled_paragraph(doc, 'Phương pháp: Cross-Sectional Z-Score Ranking | Mô hình: HMM, GARCH, Monte Carlo, Ensemble Forecast',
                          size=9, color=CLR_GRAY, italic=True)
    add_styled_paragraph(doc,
        'Cảnh báo: Báo cáo này chỉ mang tính chất tham khảo. Quyết định đầu tư cuối cùng thuộc về '
        'nhà đầu tư và cần kết hợp với phân tích cơ bản, điều kiện thị trường thực tế và khẩu vị rủi ro cá nhân.',
        size=9, color=CLR_RED, italic=True)

    # Save
    doc.save(output_path)
    log.info(f"DOCX saved: {output_path}")
    return output_path


if __name__ == '__main__':
    excel = sys.argv[1] if len(sys.argv) > 1 else 'quant_report.xlsx'
    output = sys.argv[2] if len(sys.argv) > 2 else 'bao_cao_quant.docx'
    top_n = int(sys.argv[3]) if len(sys.argv) > 3 else 5
    generate_report(excel, output, top_n)
