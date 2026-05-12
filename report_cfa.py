"""
report_cfa.py — Vietnamese Word report cho CFA-grade valuation
================================================================
5-paragraph template per ticker:
    §1. Sector & Quality Snapshot
    §2. Cost of Equity Derivation
    §3. Three-Method Valuation
    §4. Sensitivity & Risk
    §5. Conclusion & MOS

Khớp style với report.py cũ trong quant_pipeline_v4.x.
"""

from __future__ import annotations

import logging
import os
from datetime import datetime
from io import BytesIO
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from valuation_cfa import CFG, ValuationResult

log = logging.getLogger("report_cfa")

# ============================================================
# STYLING CONSTANTS
# ============================================================

COLOR_GREEN  = RGBColor(0x2E, 0x86, 0x4B)
COLOR_RED    = RGBColor(0xC0, 0x39, 0x2B)
COLOR_AMBER  = RGBColor(0xD4, 0x8C, 0x1F)
COLOR_BLUE   = RGBColor(0x2C, 0x3E, 0x70)
COLOR_GRAY   = RGBColor(0x55, 0x55, 0x55)

FONT_DEFAULT = "Calibri"

METHOD_NAME_VI = {
    "FCFF": "DCF (FCFF 2-stage)",
    "RI":   "Residual Income (EBO)",
    "PE":   "P/E Multiple",
    "PB":   "P/B Multiple",
    "EV_EBITDA": "EV/EBITDA",
}


# ============================================================
# HELPER FUNCTIONS
# ============================================================

def _set_font(run, name=FONT_DEFAULT, size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _add_heading(doc, text, level=1, color=COLOR_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    size = {1: 16, 2: 13, 3: 11}.get(level, 11)
    _set_font(run, size=size, bold=True, color=color)
    return p


def _add_para(doc, text, size=11, bold=False, color=None, align=None, indent_cm=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    if align:
        p.alignment = align
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold, color=color)
    return p


def _fmt_vnd(x: float, suffix=" VND") -> str:
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return "N/A"
    return f"{x:,.0f}{suffix}"


def _fmt_pct(x: float) -> str:
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return "N/A"
    return f"{x:+.1%}" if abs(x) < 10 else f"{x:.1%}"


def _verdict_color(verdict: str) -> RGBColor:
    if "HẤP DẪN" in verdict or "🟢" in verdict:    return COLOR_GREEN
    if "CAO" in verdict     or "🔴" in verdict:    return COLOR_RED
    return COLOR_AMBER


# ============================================================
# CHART: SENSITIVITY TORNADO (matplotlib → embedded)
# ============================================================

def _sensitivity_chart(r: ValuationResult) -> BytesIO | None:
    """Bar chart 2D: g × WACC heatmap, return PNG buffer."""
    if r.sensitivity is None or r.sensitivity.empty:
        return None
    df = r.sensitivity.copy().set_index("g")

    fig, ax = plt.subplots(figsize=(7, 3.5))
    data = df.values.astype(float)

    im = ax.imshow(data, cmap="RdYlGn", aspect="auto")
    ax.set_xticks(range(len(df.columns)))
    ax.set_xticklabels(df.columns, rotation=0, fontsize=9)
    ax.set_yticks(range(len(df.index)))
    ax.set_yticklabels(df.index, fontsize=9)
    ax.set_xlabel("Δ WACC", fontsize=10)
    ax.set_ylabel("g terminal", fontsize=10)
    ax.set_title(f"Sensitivity: Fair Value FCFF ({r.ticker})", fontsize=11, fontweight="bold")

    # Annotate cells
    for i in range(data.shape[0]):
        for j in range(data.shape[1]):
            val = data[i, j]
            if not np.isnan(val):
                ax.text(j, i, f"{val:,.0f}", ha="center", va="center",
                        fontsize=8, color="black")

    plt.colorbar(im, ax=ax, label="Fair value (VND)")
    plt.tight_layout()

    buf = BytesIO()
    plt.savefig(buf, format="png", dpi=130, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _method_comparison_chart(r: ValuationResult) -> BytesIO | None:
    """Bar chart: fair value mỗi method vs market price."""
    success_methods = {k: v for k, v in r.methods.items() if v.success}
    if not success_methods:
        return None

    labels = [METHOD_NAME_VI.get(k, k) for k in success_methods.keys()]
    fair_values = [v.fair_value for v in success_methods.values()]
    weights = [r.weights_applied.get(k, 0) for k in success_methods.keys()]

    fig, ax = plt.subplots(figsize=(7, 3.5))
    bars = ax.barh(labels, fair_values, color=COLOR_BLUE.rgb if hasattr(COLOR_BLUE, "rgb") else "#2C3E70")
    ax.axvline(r.market_price, color="red", linestyle="--", linewidth=2,
               label=f"Giá thị trường: {r.market_price:,.0f}")
    ax.axvline(r.fair_value, color="green", linestyle="-", linewidth=2,
               label=f"Fair value (weighted): {r.fair_value:,.0f}")

    for bar, fv, w in zip(bars, fair_values, weights):
        ax.text(fv, bar.get_y() + bar.get_height()/2,
                f" {fv:,.0f} (w={w:.0%})",
                va="center", fontsize=9)

    ax.set_xlabel("Fair value per share (VND)", fontsize=10)
    ax.set_title(f"So sánh phương pháp định giá — {r.ticker}", fontsize=11, fontweight="bold")
    ax.legend(loc="lower right", fontsize=8)
    plt.tight_layout()

    buf = BytesIO()
    plt.savefig(buf, format="png", dpi=130, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# ============================================================
# 5-PARAGRAPH TEMPLATE PER TICKER
# ============================================================

def _write_ticker_section(doc: Document, r: ValuationResult):
    # === Header với verdict ===
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run(f"▌{r.ticker}")
    _set_font(run, size=18, bold=True, color=COLOR_BLUE)
    run2 = p.add_run(f"   ({r.sector})")
    _set_font(run2, size=12, color=COLOR_GRAY)

    p = doc.add_paragraph()
    run = p.add_run(f"Verdict: {r.verdict}")
    _set_font(run, size=14, bold=True, color=_verdict_color(r.verdict))
    run2 = p.add_run(f"     MOS: {_fmt_pct(r.mos)}     ")
    _set_font(run2, size=12, bold=True, color=_verdict_color(r.verdict))
    run3 = p.add_run(f"Market: {_fmt_vnd(r.market_price)}  →  Fair: {_fmt_vnd(r.fair_value)}")
    _set_font(run3, size=11, color=COLOR_GRAY)

    if r.status == "FAILED":
        _add_para(doc, f"❌ KHÔNG ĐỊNH GIÁ ĐƯỢC — lý do: {r.fail_reason}",
                  size=11, bold=True, color=COLOR_RED)
        return

    # === §1. Sector & Quality Snapshot ===
    _add_heading(doc, "§1. Hồ sơ ngành & Chất lượng doanh nghiệp", level=2)
    sector_note = {
        "Banking":      "ngân hàng — ưu tiên Residual Income do tính chất book-value-heavy "
                        "và mô hình kinh doanh chênh lệch lãi suất.",
        "Insurance":    "bảo hiểm — tương tự ngân hàng, RI và P/B là primary.",
        "Real_Estate":  "bất động sản — FCFF kết hợp NAV, P/B làm sanity check.",
        "Securities":   "chứng khoán — P/B và RI phù hợp do tài sản chủ yếu là financial assets.",
        "Industrial":   "công nghiệp — FCFF 2-stage là primary, P/E sanity check.",
        "Consumer":     "tiêu dùng — FCFF + P/E (smooth EPS) làm trọng số chính.",
        "Utilities":    "tiện ích — dòng tiền ổn định, FCFF với g_terminal cao.",
        "Tech":         "công nghệ — FCFF tăng trưởng cao, P/E elevated.",
        "Materials":    "vật liệu cơ bản — FCFF với cyclical adjustment.",
        "Healthcare":   "dược/y tế — FCFF + P/E, defensive sector.",
        "default":      "ngành chưa phân loại — dùng default weights.",
    }.get(r.sector, "ngành chưa phân loại.")

    flags_text = "; ".join(r.quality_flags) if r.quality_flags else "Không có cảnh báo chất lượng."
    para1 = (
        f"{r.ticker} thuộc {sector_note} Đánh giá chất lượng doanh nghiệp: {flags_text} "
        f"Beta rolling 2Y so với VN-Index: {r.beta:.2f}, phản ánh "
        f"{'rủi ro hệ thống thấp hơn thị trường' if r.beta < 1 else 'rủi ro hệ thống cao hơn thị trường' if r.beta > 1 else 'rủi ro tương đương thị trường'}."
    )
    _add_para(doc, para1, size=11)

    # === §2. Cost of Equity ===
    _add_heading(doc, "§2. Chi phí vốn chủ sở hữu (CAPM)", level=2)
    para2 = (
        f"Áp dụng CAPM với các tham số macro Việt Nam: rf = {CFG['rf']:.1%} (TPCP 10Y), "
        f"ERP_VN = {CFG['erp_vn']:.1%} (Damodaran frontier market 2026), β = {r.beta:.2f}. "
        f"Cost of equity r_e = rf + β × ERP = {CFG['rf']:.1%} + {r.beta:.2f} × {CFG['erp_vn']:.1%} "
        f"= {r.cost_of_equity:.2%}. "
    )

    # Add WACC if FCFF was used
    fcff_result = r.methods.get("FCFF")
    if fcff_result and fcff_result.success:
        wacc = fcff_result.details.get("wacc", r.cost_of_equity)
        para2 += (
            f"WACC điều chỉnh theo cấu trúc vốn = {wacc:.2%}, dùng cho discount FCFF. "
            f"Khoảng cách WACC vs g_terminal ({CFG['g_terminal']:.1%}): "
            f"{(wacc - CFG['g_terminal']):.2%} — "
            f"{'an toàn cho terminal value' if (wacc - CFG['g_terminal']) > 0.03 else 'mong manh, cần kiểm tra sensitivity'}."
        )
    _add_para(doc, para2, size=11)

    # === §3. Three-Method Valuation ===
    _add_heading(doc, "§3. Định giá tam giác (3 phương pháp)", level=2)

    # Method table
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0].cells
    headers = ["Phương pháp", "Fair value (VND)", "Confidence", "Weight final", "Status"]
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        _set_font(run, size=10, bold=True, color=COLOR_BLUE)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for m_name, m in r.methods.items():
        row = table.add_row().cells
        row[0].text = METHOD_NAME_VI.get(m_name, m_name)
        if m.success:
            row[1].text = f"{m.fair_value:,.0f}"
            row[2].text = f"{m.confidence:.2f}"
            row[3].text = f"{r.weights_applied.get(m_name, 0):.0%}"
            row[4].text = "✓ SUCCESS"
        else:
            row[1].text = "—"
            row[2].text = "—"
            row[3].text = "—"
            row[4].text = f"✗ {m.error}"
        for c in row:
            for p in c.paragraphs:
                for run in p.runs:
                    _set_font(run, size=9)

    # Chart so sánh
    chart_buf = _method_comparison_chart(r)
    if chart_buf:
        doc.add_picture(chart_buf, width=Cm(16))

    # Narrative
    success_methods = [k for k, m in r.methods.items() if m.success]
    para3 = (
        f"Áp dụng {len(success_methods)} phương pháp ({', '.join(METHOD_NAME_VI.get(k, k) for k in success_methods)}). "
        f"Triangulation theo confidence-weighted: fair value tổng hợp = {_fmt_vnd(r.fair_value)}. "
    )
    if "FCFF" in r.weights_applied:
        para3 += f"FCFF chiếm tỷ trọng {r.weights_applied['FCFF']:.0%} do dòng tiền tự do là driver chính. "
    if "RI" in r.weights_applied:
        para3 += f"Residual Income chiếm {r.weights_applied['RI']:.0%}, phản ánh excess return so với cost of equity. "
    _add_para(doc, para3, size=11)

    # === §4. Sensitivity & Risk ===
    _add_heading(doc, "§4. Phân tích độ nhạy & Rủi ro", level=2)
    sens_buf = _sensitivity_chart(r)
    if sens_buf:
        doc.add_picture(sens_buf, width=Cm(16))
        para4 = (
            f"Grid sensitivity: g_terminal ∈ [2%, 5%] × WACC ± 1%. "
            f"Fair value FCFF biến động trong dải này cho thấy mức độ nhạy của định giá. "
            f"Downside scenario (g = 2%, WACC + 1%) cần được nhà đầu tư cân nhắc khi entry với position size lớn."
        )
    else:
        para4 = (
            "Mã không thực hiện sensitivity FCFF (sector tài chính hoặc thiếu data dòng tiền). "
            "Rủi ro chính nằm ở giả định ROE bền vững (RI) và P/E sector benchmark — "
            "cần monitor BCTC quarterly và re-rate khi sector benchmark thay đổi."
        )
    _add_para(doc, para4, size=11)

    # === §5. Conclusion & Action ===
    _add_heading(doc, "§5. Kết luận & Hành động", level=2)
    if "HẤP DẪN" in r.verdict:
        action = (
            f"MOS = {_fmt_pct(r.mos)} vượt ngưỡng +20% — cổ phiếu đang giao dịch dưới fair value đáng kể. "
            f"Khuyến nghị: kết hợp với tín hiệu kỹ thuật BUY từ quant_pipeline để xác định entry "
            f"(regime BULL hoặc đầu BEAR-rebound). Position size đề xuất theo Mean-CVaR optimizer. "
            f"Stop-loss: cắt nếu giá đóng cửa dưới -8% từ entry hoặc khi narrative cơ bản thay đổi."
        )
    elif "CAO" in r.verdict:
        action = (
            f"MOS = {_fmt_pct(r.mos)} dưới -10% — giá thị trường đang vượt fair value đáng kể. "
            f"Khuyến nghị: KHÔNG entry mới dù có tín hiệu kỹ thuật BUY. "
            f"Nếu đang nắm giữ, cân nhắc trim position hoặc đặt trailing stop chặt. "
            f"Đợi pullback về vùng gần fair value hoặc re-rate khi có catalyst tăng trưởng mới."
        )
    else:
        action = (
            f"MOS = {_fmt_pct(r.mos)} trong vùng fair value (-10% đến +20%). "
            f"Cổ phiếu đang được định giá hợp lý — không có margin of safety đủ lớn để aggressive long. "
            f"Có thể giao dịch theo tín hiệu kỹ thuật ngắn hạn với position size nhỏ và "
            f"stop-loss chặt. Hold nếu đã có vị thế, chờ catalyst để re-evaluate."
        )
    _add_para(doc, action, size=11)

    # Page break
    doc.add_page_break()


# ============================================================
# OVERVIEW SECTION
# ============================================================

def _write_overview(doc: Document, results: list[ValuationResult]):
    _add_heading(doc, "TỔNG QUAN", level=1)

    n_total = len(results)
    n_succ  = sum(1 for r in results if r.status == "SUCCESS")
    n_part  = sum(1 for r in results if r.status == "PARTIAL")
    n_fail  = sum(1 for r in results if r.status == "FAILED")
    n_green = sum(1 for r in results if "HẤP DẪN" in r.verdict)
    n_amber = sum(1 for r in results if "FAIR" in r.verdict)
    n_red   = sum(1 for r in results if "CAO" in r.verdict)

    valid_mos = [r.mos for r in results if not np.isnan(r.mos)]
    median_mos = np.median(valid_mos) if valid_mos else np.nan

    table = doc.add_table(rows=6, cols=2)
    table.style = "Light Grid Accent 1"
    rows_data = [
        ("Tổng số mã đầu vào", str(n_total)),
        ("Định giá SUCCESS (≥2 methods)", f"{n_succ} ({n_succ/n_total:.0%})"),
        ("Định giá PARTIAL (1 method)", f"{n_part} ({n_part/n_total:.0%})" if n_total else "0"),
        ("🟢 Hấp dẫn (MOS > +20%)", str(n_green)),
        ("🟡 Gần fair value", str(n_amber)),
        ("🔴 Định giá cao (MOS ≤ -10%)", str(n_red)),
    ]
    for i, (k, v) in enumerate(rows_data):
        c = table.rows[i].cells
        c[0].text = k
        c[1].text = v
        for run in c[0].paragraphs[0].runs:
            _set_font(run, size=11, bold=True)
        for run in c[1].paragraphs[0].runs:
            _set_font(run, size=11)

    _add_para(doc, "")
    _add_para(doc, f"MOS trung vị: {_fmt_pct(median_mos)}", size=12, bold=True, color=COLOR_BLUE)

    # === Bảng tổng hợp ===
    _add_heading(doc, "Bảng tổng hợp định giá", level=2)
    if n_succ + n_part == 0:
        _add_para(doc, "Không có mã nào định giá thành công.", size=11, color=COLOR_RED)
        return

    success_results = [r for r in results if r.status in ("SUCCESS", "PARTIAL")]
    # Sort by MOS desc
    success_results.sort(key=lambda r: r.mos if not np.isnan(r.mos) else -999, reverse=True)

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Ticker", "Sector", "Market", "Fair value", "MOS", "Verdict"]):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        _set_font(run, size=10, bold=True, color=COLOR_BLUE)

    for r in success_results:
        row = table.add_row().cells
        row[0].text = r.ticker
        row[1].text = r.sector
        row[2].text = _fmt_vnd(r.market_price, "")
        row[3].text = _fmt_vnd(r.fair_value, "")
        row[4].text = _fmt_pct(r.mos)
        row[5].text = r.verdict
        for c in row:
            for p in c.paragraphs:
                for run in p.runs:
                    _set_font(run, size=10)

    # Bảng FAILED
    failed = [r for r in results if r.status == "FAILED"]
    if failed:
        _add_heading(doc, "Mã không định giá được", level=2)
        for r in failed:
            _add_para(doc, f"• {r.ticker} — {r.fail_reason}", size=10, color=COLOR_RED)

    doc.add_page_break()


# ============================================================
# METHODOLOGY APPENDIX
# ============================================================

def _write_methodology(doc: Document):
    _add_heading(doc, "PHỤ LỤC: PHƯƠNG PHÁP LUẬN CFA", level=1)

    _add_heading(doc, "1. DCF — Free Cash Flow to Firm (2-stage)", level=2)
    _add_para(doc,
        "FCFF = Net Income + D&A + Interest×(1-t) - CapEx - ΔWC. "
        "Discount với WACC = (E/V)×r_e + (D/V)×r_d×(1-t). "
        "Stage 1: 5 năm tăng trưởng linear fade từ g_3Y → g_terminal. "
        "Stage 2: Gordon Growth Model với g_terminal = 4% (~ inflation dài hạn). "
        "Áp dụng: industrial, consumer, tech, utilities, real estate. KHÔNG dùng cho banking/insurance.")

    _add_heading(doc, "2. Residual Income — Edwards-Bell-Ohlson", level=2)
    _add_para(doc,
        "V₀ = BV₀ + (ROE - r) × BV₀ / (1 + r - ω), với ω = 0.6 (persistence factor). "
        "ROE bền vững = trung bình 3 năm gần nhất. "
        "Khi ROE ≤ r → fair value = book value (không có excess return). "
        "Áp dụng tốt nhất: banking, insurance, securities, BV-heavy stocks. "
        "Robust hơn DCF cho thị trường Việt Nam do BCTC banking khó forecast FCFF.")

    _add_heading(doc, "3. Relative Valuation — Sector Multiples", level=2)
    _add_para(doc,
        "P/E: Fair = EPS_smooth × P/E_target_sector. EPS smooth = trung bình 3 năm. "
        "P/B: Fair = BVPS × P/B_target_sector. Áp dụng cho banking, securities. "
        "EV/EBITDA: dùng cho capital-intensive (utilities, real estate). "
        "Vai trò: sanity check, không phải primary method. "
        "Multiples refresh quarterly theo benchmark sector VN.")

    _add_heading(doc, "4. Triangulation (Confidence-Weighted)", level=2)
    _add_para(doc,
        "Weight final mỗi method = base_weight_sector × confidence_score, sau đó normalize. "
        "Confidence FCFF cao khi: ≥5 năm FCF dương, growth ổn định, WACC - g > 3%. "
        "Confidence RI cao khi: ROE 3Y ổn định, excess return > 2%, sector phù hợp. "
        "Confidence Multiples baseline 0.7, boost cho sector phù hợp (P/B cho banking).")

    _add_heading(doc, "5. Tham số Macro VN (CFG)", level=2)
    for k, v in [
        ("Risk-free rate (TPCP 10Y)", f"{CFG['rf']:.1%}"),
        ("ERP Việt Nam (Damodaran 2026)", f"{CFG['erp_vn']:.1%}"),
        ("Terminal growth", f"{CFG['g_terminal']:.1%}"),
        ("Tax rate (CIT VN)", f"{CFG['tax_rate']:.0%}"),
        ("Default beta", f"{CFG['default_beta']:.2f}"),
        ("Stage 1 horizon (years)", f"{CFG['high_growth_years']}"),
    ]:
        _add_para(doc, f"  • {k}: {v}", size=10)

    _add_heading(doc, "6. Ngưỡng phân loại Verdict", level=2)
    _add_para(doc, f"  • 🟢 Hấp dẫn: MOS > +{CFG['mos_attractive']:.0%}", size=10, color=COLOR_GREEN)
    _add_para(doc, f"  • 🟡 Gần fair value: {CFG['mos_expensive']:.0%} < MOS ≤ +{CFG['mos_attractive']:.0%}",
              size=10, color=COLOR_AMBER)
    _add_para(doc, f"  • 🔴 Định giá cao: MOS ≤ {CFG['mos_expensive']:.0%}", size=10, color=COLOR_RED)

    _add_para(doc, "")
    _add_para(doc,
        "⚠️ Disclaimer: Báo cáo này được tạo tự động từ Naizy Quant Pipeline — Valuation Layer. "
        "Các giả định macro (rf, ERP, g) cần được review định kỳ theo điều kiện thị trường. "
        "Định giá nội tại KHÔNG phải khuyến nghị đầu tư — cần kết hợp với tín hiệu kỹ thuật, "
        "regime detection và risk management từ các module khác của pipeline.",
        size=10, color=COLOR_GRAY)


# ============================================================
# MAIN: BUILD REPORT
# ============================================================

def build_word_report(
    results: list[ValuationResult],
    output_path: str | Path | None = None,
) -> Path:
    """Build Word report tiếng Việt từ list ValuationResult."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_DEFAULT
    style.font.size = Pt(11)

    # === Title page ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("BÁO CÁO ĐỊNH GIÁ NỘI TẠI")
    _set_font(run, size=22, bold=True, color=COLOR_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Naizy Quant Pipeline — Valuation Layer (CFA L2 Methodology)")
    _set_font(run, size=12, color=COLOR_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    run = p.add_run(f"Ngày phát hành: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    _set_font(run, size=11, color=COLOR_GRAY)

    # === Overview ===
    _write_overview(doc, results)

    # === Detail per ticker ===
    _add_heading(doc, "CHI TIẾT TỪNG MÃ", level=1)
    for r in results:
        if r.status in ("SUCCESS", "PARTIAL"):
            _write_ticker_section(doc, r)

    # === Failed tickers (brief) ===
    failed = [r for r in results if r.status == "FAILED"]
    if failed:
        _add_heading(doc, "CHI TIẾT MÃ KHÔNG ĐỊNH GIÁ ĐƯỢC", level=1)
        for r in failed:
            _write_ticker_section(doc, r)

    # === Methodology ===
    _write_methodology(doc)

    # Save
    if output_path is None:
        ts = datetime.now().strftime("%Y%m%d_%H%M")
        out_dir = Path(CFG.get("output_dir", "./output"))
        out_dir.mkdir(parents=True, exist_ok=True)
        output_path = out_dir / f"valuation_report_{ts}.docx"
    output_path = Path(output_path)

    doc.save(output_path)
    log.info(f"Report saved: {output_path}")
    return output_path
