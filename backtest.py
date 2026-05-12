"""
backtest.py
===========
All-in-one forecast backtesting pipeline.

Reads quant_report_gen.py .docx outputs, extracts recommendations,
evaluates them against OHLC data using triple-barrier labeling
(López de Prado, AFML Ch.3), and writes summary CSVs.

USAGE
-----
1. Place .docx reports in any folder (e.g. ./reports/ or current dir)
2. Set REPORTS_DIR below, or pass via CLI: python backtest.py ./reports
3. To use real OHLC: replace synthetic_price_fn with vnstock fetch
   (see PRODUCTION SECTION at bottom)

OUTPUTS (in ./bt_output/)
-------
01_recommendations.csv  : parsed recs from .docx
02_evaluated.csv        : recs + realized outcomes
03_per_ticker.csv       : compact per-trade evaluation
04_calibration.csv      : reliability diagram by Conf bin
"""

from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Optional

import numpy as np
import pandas as pd


# =============================================================================
# CONFIGURATION
# =============================================================================

REPORTS_DIR = Path("./reports")     # change as needed
OUTPUT_DIR  = Path("./bt_output")


@dataclass
class EvalConfig:
    """All evaluation knobs in one place."""
    entry_tolerance_pct: float = 0.02     # entry "touched" if Low <= entry*(1+tol)
    entry_window_days: int = 5            # must reach entry within N sessions
    horizons: tuple = (3, 5, 10, 20)      # fixed-horizon return snapshots
    vertical_barrier_days: int = 20       # time-stop for triple barrier
    cost_bps: float = 40.0                # round-trip cost (VN ~0.40%)


# =============================================================================
# SECTION 1: EXTRACTOR — parse .docx into normalized DataFrame
# =============================================================================

# ---------------------------------------------------------------------------
# SECTION 1 — EXTRACTOR  (reads python-docx objects directly, no CLI tools)
# ---------------------------------------------------------------------------
# Document structure confirmed from actual .docx:
#   PHẦN II table  → header ['#','Mã','Score','Rating','Timing','Forecast',
#                             'Conf','Entry','SL','TP2R','TP3R']
#   Per-ticker     → header ['Nhóm','Chỉ số','Giá trị','Đánh giá']
#                    follows a Heading 2 like "1. VIC — Điểm 89 | SWING BUY"
# ---------------------------------------------------------------------------

def _open_docx(docx_path: Path):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Run: pip install python-docx") from exc
    return Document(docx_path)


def _parse_report_date(docx_path: Path) -> pd.Timestamp:
    """Pull date from filename.
    Supports: DD-MM-YY(YY), and DD-MM (uses current year as fallback).
    """
    stem = docx_path.stem
    m = re.search(r"(\d{1,2})-(\d{2})-(\d{2,4})", stem)
    if m:
        dd, mm, yy = m.groups()
        yy_full = int(yy) if len(yy) == 4 else 2000 + int(yy)
        return pd.Timestamp(year=yy_full, month=int(mm), day=int(dd))
    m2 = re.search(r"(\d{1,2})-(\d{2})", stem)
    if m2:
        dd, mm = m2.groups()
        return pd.Timestamp(year=pd.Timestamp.now().year, month=int(mm), day=int(dd))
    raise ValueError(f"Cannot parse date from: {stem}")


def _table_header(table) -> list:
    if not table.rows:
        return []
    return [c.text.strip() for c in table.rows[0].cells]


def _table_to_dict_rows(table) -> list:
    headers = _table_header(table)
    rows = []
    for row in table.rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        rows.append(dict(zip(headers, cells)))
    return rows


def _parse_top5_table_from_doc(doc) -> pd.DataFrame:
    """Find and parse the PHẦN II top-5 table directly from docx tables."""
    target_header = {"#", "Mã", "Score", "Rating", "Forecast",
                     "Conf", "Entry", "SL", "TP2R", "TP3R"}
    for table in doc.tables:
        if target_header.issubset(set(_table_header(table))):
            rows = _table_to_dict_rows(table)
            parsed = []
            for r in rows:
                try:
                    ens = float(r.get("Forecast","0").replace("%","").replace("+",""))
                    conf = int(r.get("Conf","0").replace("%",""))
                    parsed.append({
                        "rank":         int(r.get("#", 0)),
                        "ticker":       r.get("Mã","").strip(),
                        "score":        int(r.get("Score", 0)),
                        "rating":       r.get("Rating","").strip(),
                        "timing_short": r.get("Timing","").strip(),
                        "ens_ret_pct":  ens,
                        "conf":         conf,
                        "entry":        float(r.get("Entry", 0)),
                        "sl":           float(r.get("SL", 0)),
                        "tp2r":         float(r.get("TP2R", 0)),
                        "tp3r":         float(r.get("TP3R", 0)),
                    })
                except (ValueError, KeyError):
                    continue
            if parsed:
                return pd.DataFrame(parsed)
    return pd.DataFrame()


def _parse_ticker_details_from_doc(doc, ticker: str) -> dict:
    """Parse per-ticker detail table for a given ticker symbol.

    Walks body elements: waits for a Heading 2 containing the ticker,
    then reads the next table with header [Nhóm|Chỉ số|Giá trị|Đánh giá].
    """
    try:
        from docx.text.paragraph import Paragraph
        from docx.table import Table as DocxTable
    except ImportError:
        return {}

    found_heading = False
    for child in doc.element.body:
        tag = child.tag.split("}")[-1]
        if tag == "p":
            para = Paragraph(child, doc)
            if para.style and "Heading 2" in para.style.name:
                if re.search(rf"\b{re.escape(ticker)}\b", para.text):
                    found_heading = True
                else:
                    if found_heading:
                        found_heading = False  # next ticker section
        elif tag == "tbl" and found_heading:
            table = DocxTable(child, doc)
            hdr = _table_header(table)
            if "Chỉ số" in hdr and "Giá trị" in hdr:
                result = {}
                for row in _table_to_dict_rows(table):
                    metric = row.get("Chỉ số", "").strip()
                    val    = row.get("Giá trị", "").strip()
                    if not metric or not val:
                        continue
                    val_clean = val.replace("%", "").replace("+", "")
                    try:
                        val_f = float(val_clean)
                    except ValueError:
                        val_f = None
                    if metric == "AnnRet":
                        result["ann_ret"] = val_f
                    elif metric == "Sharpe":
                        result["sharpe"] = val_f
                    elif metric == "WinRate":
                        result["winrate"] = val_f
                    elif metric == "MaxDD":
                        result["maxdd"] = val_f
                    elif metric == "HMM":
                        result["hmm"] = val
                    elif metric == "VolReg":
                        result["volreg"] = val
                    elif metric == "Forecast":
                        result["forecast_dir"] = val
                    elif metric == "LockDD%":
                        result["lock_dd_pct"] = val_f
                    elif metric == "PLoss3%":
                        result["ploss3_pct"] = val_f
                return result
    return {}


def extract_one(docx_path: Path) -> pd.DataFrame:
    """Parse one .docx into a tidy DataFrame (1 row per ticker)."""
    doc = _open_docx(docx_path)
    report_date = _parse_report_date(docx_path)

    vni = "UNKNOWN"
    for para in doc.paragraphs[:10]:
        m = re.search(r"Bối cảnh VNI:\s*([A-Z]+)", para.text)
        if m:
            vni = m.group(1)
            break

    top5 = _parse_top5_table_from_doc(doc)
    if top5.empty:
        return pd.DataFrame()

    detail_df = pd.DataFrame(
        [_parse_ticker_details_from_doc(doc, t) for t in top5.ticker]
    )
    out = pd.concat([top5.reset_index(drop=True), detail_df], axis=1)
    out.insert(0, "report_date", report_date)
    out.insert(1, "vni_context", vni)
    return out


def _extract_one_xlsx(xlsx_path: Path) -> pd.DataFrame:
    """Extract recommendations from khuyến_nghị_*.xlsx files.

    Expects sheet 'Danh Sách Mua' with header row at row index 2
    (0-based), columns: Mã, Score, Rating, HMM, Forecast, EnsRet%,
    Conf%, Entry, SL, TP1(2R), TP2(3R), Timing.
    Date is parsed from filename DD-MM-YY.
    """
    report_date = _parse_report_date(xlsx_path)

    try:
        raw = pd.read_excel(xlsx_path, sheet_name="Danh Sách Mua", header=None)
    except Exception as exc:
        raise ValueError(f"Cannot open sheet 'Danh Sách Mua': {exc}")

    # Find header row: the row that contains "Mã" and "Score"
    header_row = None
    for i, row in raw.iterrows():
        vals = [str(v).strip() for v in row.values]
        if "Mã" in vals and "Score" in vals:
            header_row = i
            break
    if header_row is None:
        raise ValueError("Cannot locate header row with 'Mã' and 'Score'")

    df = pd.read_excel(xlsx_path, sheet_name="Danh Sách Mua",
                       header=header_row)
    # Drop fully-empty rows and the leading NaN column
    df = df.dropna(how="all").reset_index(drop=True)
    # Keep only rows where 'Mã' looks like a ticker (2-5 uppercase letters)
    if "Mã" not in df.columns:
        raise ValueError(f"Column 'Mã' not found. Columns: {list(df.columns)}")
    df = df[df["Mã"].astype(str).str.match(r"^[A-Z]{2,5}$")].reset_index(drop=True)
    if df.empty:
        return pd.DataFrame()

    # Clean emoji/unicode from string columns
    import re as _re
    def _clean(s):
        if pd.isna(s):
            return ""
        return _re.sub(r"[^\w\s\-+.%:/()&]", "", str(s)).strip()

    def _parse_float(s):
        try:
            return float(_re.sub(r"[^\d.\-+]", "", str(s)))
        except ValueError:
            return np.nan

    # Map columns (handle both TP1(2R) and TP2R naming)
    col_tp2 = "TP1(2R)" if "TP1(2R)" in df.columns else "TP2R"
    col_tp3 = "TP2(3R)" if "TP2(3R)" in df.columns else "TP3R"

    rows = []
    for i, r in df.iterrows():
        hmm_raw   = _clean(r.get("HMM", ""))
        hmm_clean = _re.sub(r"\s+", " ", hmm_raw).split()[-1] if hmm_raw else "UNKNOWN"
        rows.append({
            "rank":          i + 1,
            "ticker":        str(r["Mã"]).strip(),
            "score":         int(_parse_float(r.get("Score", 0)) or 0),
            "rating":        _clean(r.get("Rating", "")),
            "timing_short":  _clean(r.get("Timing", "")),
            "ens_ret_pct":   _parse_float(r.get("EnsRet%", 0)),
            "conf":          int((_parse_float(r.get("Conf%", 0)) or 0) * 100)
                             if _parse_float(r.get("Conf%", 0) or 0) <= 1
                             else int(_parse_float(r.get("Conf%", 0)) or 0),
            "entry":         _parse_float(r.get("Entry", 0)),
            "sl":            _parse_float(r.get("SL", 0)),
            "tp2r":          _parse_float(r.get(col_tp2, 0)),
            "tp3r":          _parse_float(r.get(col_tp3, 0)),
            "hmm":           hmm_clean,
            "volreg":        None,
            "forecast_dir":  "TĂNG",
            "ann_ret":       np.nan,
            "sharpe":        np.nan,
            "winrate":       np.nan,
            "maxdd":         np.nan,
            "lock_dd_pct":   np.nan,
            "ploss3_pct":    np.nan,
        })

    out = pd.DataFrame(rows)
    out.insert(0, "report_date", report_date)

    # VNI context from 'Tổng Quan Mua' sheet if available
    vni = "UNKNOWN"
    try:
        meta = pd.read_excel(xlsx_path, sheet_name="Tổng Quan Mua", header=None)
        for _, row in meta.iterrows():
            for cell in row.values:
                m = _re.search(r"VNI:\s*([A-Z]+)", str(cell))
                if m:
                    vni = m.group(1)
                    break
            if vni != "UNKNOWN":
                break
    except Exception:
        pass
    out.insert(1, "vni_context", vni)
    return out


def extract_one(docx_path: Path) -> pd.DataFrame:
    """Parse one .docx into a tidy DataFrame (1 row per ticker)."""
    doc = _open_docx(docx_path)
    report_date = _parse_report_date(docx_path)

    vni = "UNKNOWN"
    for para in doc.paragraphs[:10]:
        m = re.search(r"Bối cảnh VNI:\s*([A-Z]+)", para.text)
        if m:
            vni = m.group(1)
            break

    top5 = _parse_top5_table_from_doc(doc)
    if top5.empty:
        return pd.DataFrame()

    detail_df = pd.DataFrame(
        [_parse_ticker_details_from_doc(doc, t) for t in top5.ticker]
    )
    out = pd.concat([top5.reset_index(drop=True), detail_df], axis=1)
    out.insert(0, "report_date", report_date)
    out.insert(1, "vni_context", vni)
    return out


def extract_many(paths: list) -> pd.DataFrame:
    """Batch-extract from .docx or .xlsx files (auto-detected by extension).

    Pass a folder Path, a list of .docx paths, or a list of .xlsx paths.
    If a folder is passed, prefers *khuyến_nghị*.xlsx files (full list)
    over *báo_cáo*.docx (top-5 only).
    """
    # Accept a single folder → auto-discover files
    if len(paths) == 1 and Path(paths[0]).is_dir():
        folder = Path(paths[0])
        xlsx_files = sorted(folder.glob("khuyến_nghị*.xlsx")) +                      sorted(folder.glob("khuyen_nghi*.xlsx"))
        docx_files = sorted(folder.glob("báo_cáo*.docx")) +                      sorted(folder.glob("bao_cao*.docx")) +                      sorted(folder.glob("*.docx"))
        # Prefer xlsx (full list); fall back to docx
        paths = xlsx_files if xlsx_files else docx_files

    dfs = []
    for p in paths:
        p = Path(p)
        try:
            if p.suffix.lower() == ".xlsx":
                df = _extract_one_xlsx(p)
            else:
                df = extract_one(p)
            if not df.empty:
                dfs.append(df)
                print(f"[OK] {p.name}: {len(df)} recs")
        except Exception as e:
            print(f"[WARN] {p.name} failed: {e}")
    if not dfs:
        return pd.DataFrame()
    return (pd.concat(dfs, ignore_index=True)
              .sort_values(["report_date", "rank"])
              .reset_index(drop=True))


# =============================================================================
# SECTION 2: EVALUATOR — triple barrier + directional accuracy
# =============================================================================

def _label_triple_barrier(prices: pd.DataFrame, entry_price: float,
                          sl_price: float, tp_price: float,
                          cfg: EvalConfig) -> dict:
    """
    First-touch labeling per López de Prado AFML Ch.3.
    Returns hit_label: +1 (TP first), -1 (SL first), 0 (time barrier).
    Uses High for TP, Low for SL (conservative).
    """
    if prices.empty or len(prices) < 2:
        return {"hit_label": None, "days_to_hit": None, "exit_price": None}

    for i, (dt, row) in enumerate(prices.iterrows()):
        if i == 0:
            continue  # skip entry bar
        if row["Low"] <= sl_price:
            return {"hit_label": -1, "days_to_hit": i, "exit_price": sl_price}
        if row["High"] >= tp_price:
            return {"hit_label": +1, "days_to_hit": i, "exit_price": tp_price}
        if i >= cfg.vertical_barrier_days:
            return {"hit_label": 0, "days_to_hit": i, "exit_price": row["Close"]}

    last = prices.iloc[-1]
    return {"hit_label": 0, "days_to_hit": len(prices) - 1,
            "exit_price": last["Close"]}


def _realized_horizons(prices: pd.DataFrame, entry_price: float,
                       horizons: tuple) -> dict:
    """Fixed-horizon return + max adverse excursion (drawdown) within horizon."""
    out = {}
    for h in horizons:
        if len(prices) > h:
            ret = prices["Close"].iloc[h] / entry_price - 1
            mdd = prices["Low"].iloc[1:h+1].min() / entry_price - 1
            out[f"ret_{h}d"] = ret
            out[f"mdd_{h}d"] = mdd
        else:
            out[f"ret_{h}d"] = np.nan
            out[f"mdd_{h}d"] = np.nan
    return out


def evaluate_recommendation(rec: pd.Series, price_df: pd.DataFrame,
                            cfg: EvalConfig) -> dict:
    """Evaluate one recommendation row vs OHLC. Core logic of the evaluator."""
    post = price_df[price_df.index > rec.report_date].head(cfg.entry_window_days)
    tol = rec.entry * (1 + cfg.entry_tolerance_pct)
    touched = post[post["Low"] <= tol]

    result = {"entered": False, "entry_date": None, "entry_fill": np.nan,
              "hit_label": None, "days_to_hit": None, "exit_price": np.nan,
              "trade_ret_net": np.nan, "direction_correct": None}
    for h in cfg.horizons:
        result[f"ret_{h}d"] = np.nan
        result[f"mdd_{h}d"] = np.nan

    if touched.empty:
        # Not entered — still compute realized returns for directional analysis
        if not post.empty:
            ref = post["Close"].iloc[0]
            result.update(_realized_horizons(post, ref, cfg.horizons))
            ret20 = result.get("ret_20d")
            if pd.notna(ret20):
                result["direction_correct"] = bool(
                    np.sign(ret20) == np.sign(rec.ens_ret_pct))
        return result

    entry_bar = touched.iloc[0]
    fill = min(entry_bar["Open"], rec.entry)
    result.update({"entered": True, "entry_date": entry_bar.name,
                   "entry_fill": fill})

    forward = price_df[price_df.index >= entry_bar.name]
    result.update(_label_triple_barrier(forward, fill, rec.sl, rec.tp2r, cfg))

    if (result["exit_price"] is not None
            and not np.isnan(result["exit_price"])):
        gross = result["exit_price"] / fill - 1
        result["trade_ret_net"] = gross - cfg.cost_bps / 10000

    result.update(_realized_horizons(forward, fill, cfg.horizons))

    ret20 = result.get("ret_20d")
    if pd.notna(ret20):
        result["direction_correct"] = bool(
            np.sign(ret20) == np.sign(rec.ens_ret_pct))
    return result


def evaluate_all(recs_df: pd.DataFrame,
                 price_fn: Callable[[str, pd.Timestamp, pd.Timestamp], pd.DataFrame],
                 cfg: Optional[EvalConfig] = None) -> pd.DataFrame:
    """Batch evaluate all recommendations using a price-fetching callable."""
    cfg = cfg or EvalConfig()
    results = []
    look_fwd = cfg.vertical_barrier_days + cfg.entry_window_days + 5

    for _, rec in recs_df.iterrows():
        start = rec.report_date - pd.Timedelta(days=5)
        end = rec.report_date + pd.Timedelta(days=look_fwd * 2)
        try:
            px = price_fn(rec.ticker, start, end)
        except Exception as e:
            print(f"[WARN] price fetch failed {rec.ticker}@"
                  f"{rec.report_date.date()}: {e}")
            continue
        if px is None or px.empty:
            continue
        r = evaluate_recommendation(rec, px.sort_index(), cfg)
        r["ticker"] = rec.ticker
        r["report_date"] = rec.report_date
        results.append(r)

    if not results:
        return recs_df.copy()
    out = pd.DataFrame(results)
    return recs_df.merge(out, on=["ticker", "report_date"], how="left")


# =============================================================================
# SECTION 3: AGGREGATE METRICS
# =============================================================================

def compute_summary(df: pd.DataFrame) -> dict:
    """Top-line quality metrics suitable for printing."""
    d = df.copy()
    if d.empty or "entered" not in d.columns:
        print("[WARN] compute_summary: no evaluation data (0 recommendations extracted).")
        return {"n_recs": 0}
    entered = d[d["entered"] == True]
    closed = entered[entered["hit_label"].notna()]

    def pct(x): return float(np.nan if len(x) == 0 else x.mean())

    out = {"n_recs": int(len(d)), "entry_fill_rate": pct(d["entered"])}

    # Directional accuracy at each horizon
    for h in [3, 5, 10, 20]:
        col = f"ret_{h}d"
        if col in d.columns:
            mask = d[col].notna()
            if mask.any():
                fc = np.sign(d.loc[mask, "ens_ret_pct"])
                rt = np.sign(d.loc[mask, col])
                out[f"DA_{h}d"] = float((fc == rt).mean())

    # Triple barrier breakdown
    if not closed.empty:
        out["n_trades"] = int(len(closed))
        out["tp_hit_rate"]  = pct(closed["hit_label"] == 1)
        out["sl_hit_rate"]  = pct(closed["hit_label"] == -1)
        out["time_out_rate"] = pct(closed["hit_label"] == 0)
        out["expectancy_per_trade"] = pct(closed["trade_ret_net"])
        out["median_trade_ret_net"] = float(closed["trade_ret_net"].median())
        std = closed["trade_ret_net"].std()
        out["trade_ret_std"] = float(std) if std and std > 0 else None

    # Forecast-vs-realized error (10d horizon, typical pipeline target)
    if "ret_10d" in d.columns:
        mask = d["ret_10d"].notna()
        if mask.any():
            err = (d.loc[mask, "ens_ret_pct"] / 100) - d.loc[mask, "ret_10d"]
            out["MAE_10d"]  = float(err.abs().mean())
            out["BIAS_10d"] = float(err.mean())  # +ve = over-forecasting
    return out


def per_ticker_report(df: pd.DataFrame) -> pd.DataFrame:
    """Compact per-trade table for quick inspection."""
    cols = ["report_date", "ticker", "score", "rating", "conf",
            "ens_ret_pct", "entry", "sl", "tp2r", "tp3r",
            "entered", "entry_date", "entry_fill",
            "hit_label", "days_to_hit", "trade_ret_net",
            "ret_3d", "ret_5d", "ret_10d", "ret_20d",
            "mdd_10d", "mdd_20d", "direction_correct",
            "lock_dd_pct", "ploss3_pct"]
    return df[[c for c in cols if c in df.columns]].copy()


def calibration_table(df: pd.DataFrame, by: str = "conf") -> pd.DataFrame:
    """Reliability diagram in tabular form. Compare model conf vs actual hit rate."""
    d = df[df["ret_20d"].notna()].copy()
    if d.empty:
        return pd.DataFrame()
    d["correct"] = (np.sign(d["ens_ret_pct"]) == np.sign(d["ret_20d"])).astype(int)
    tbl = d.groupby(by).agg(
        n=("correct", "size"),
        hit_rate=("correct", "mean"),
        mean_forecast_ret=("ens_ret_pct", "mean"),
        mean_realized_ret_20d=("ret_20d", lambda s: s.mean() * 100),
    ).reset_index()
    tbl["calibration_gap"] = tbl[by] - tbl["hit_rate"] * 100
    return tbl


# =============================================================================
# SECTION 4: PRICE FUNCTIONS
# =============================================================================
# Demo uses synthetic GBM. Replace with vnstock_price_fn for production.

def synthetic_price_fn(ticker: str, start: pd.Timestamp,
                       end: pd.Timestamp) -> pd.DataFrame:
    """Synthetic OHLC seeded by ticker. DEMO ONLY — not real prices."""
    seed = abs(hash(ticker)) % (2**31)
    rng = np.random.default_rng(seed)
    dates = pd.bdate_range(start, end)
    if len(dates) == 0:
        return pd.DataFrame()
    rets = rng.normal(0.0005, 0.025, size=len(dates))
    anchor = {"VHM": 150, "VIC": 200, "CDC": 19, "PVP": 18, "BAF": 37,
              "VJC": 175, "DHC": 37, "VHC": 63, "TNT": 11, "MWG": 87,
              "PTB": 50, "BFC": 66}.get(ticker, 50)
    p0 = anchor * (1 + rng.normal(0, 0.02))
    close = p0 * np.cumprod(1 + rets)
    high = close * (1 + np.abs(rng.normal(0, 0.005, size=len(dates))))
    low  = close * (1 - np.abs(rng.normal(0, 0.005, size=len(dates))))
    openp = close * (1 + rng.normal(0, 0.003, size=len(dates)))
    return pd.DataFrame({"Open": openp, "High": high, "Low": low,
                         "Close": close}, index=dates)


# --- PRODUCTION SECTION: uncomment & use this when ready ---------------------
# def vnstock_price_fn(ticker: str, start: pd.Timestamp,
#                      end: pd.Timestamp) -> pd.DataFrame:
#     """Real OHLC from vnstock (golden tier). VND scaled to k VND."""
#     from vnstock import Vnstock
#     stock = Vnstock().stock(symbol=ticker, source='VCI')
#     df = stock.quote.history(
#         start=start.strftime('%Y-%m-%d'),
#         end=end.strftime('%Y-%m-%d'),
#         interval='1D'
#     )
#     df = df.rename(columns={'time': 'Date', 'open': 'Open', 'high': 'High',
#                             'low': 'Low', 'close': 'Close'})
#     df['Date'] = pd.to_datetime(df['Date'])
#     df = df.set_index('Date').sort_index()
#     # vnstock returns VND; entry/sl/tp in reports are k VND -> scale down
#     if df['Close'].mean() > 1000:
#         df[['Open', 'High', 'Low', 'Close']] /= 1000
#     return df[['Open', 'High', 'Low', 'Close']]


# =============================================================================
# SECTION 5: ORCHESTRATOR
# =============================================================================

def run(reports_dir: Path, output_dir: Path = OUTPUT_DIR,
        price_fn: Callable = synthetic_price_fn,
        cfg: Optional[EvalConfig] = None) -> tuple:
    """Full pipeline: extract -> evaluate -> save -> print summary."""
    output_dir.mkdir(exist_ok=True)
    cfg = cfg or EvalConfig()

    # 1. extract
    print(f"Scanning {reports_dir} for xlsx/docx reports...")
    recs = extract_many([Path(reports_dir)])
    recs.to_csv(output_dir / "01_recommendations.csv", index=False)
    print(f"Extracted {len(recs)} recommendations\n")

    # 2. evaluate
    enriched = evaluate_all(recs, price_fn=price_fn, cfg=cfg)
    enriched.to_csv(output_dir / "02_evaluated.csv", index=False)

    # 3. per-ticker
    per_tkr = per_ticker_report(enriched)
    per_tkr.to_csv(output_dir / "03_per_ticker.csv", index=False)
    print("=" * 70)
    print("PER-TICKER EVALUATION")
    print("=" * 70)
    disp = per_tkr.copy()
    for c in ["ret_3d", "ret_5d", "ret_10d", "ret_20d",
              "mdd_10d", "mdd_20d", "trade_ret_net"]:
        if c in disp:
            disp[c] = (disp[c] * 100).round(2)
    print(disp.to_string(index=False))

    # 4. summary
    summary = compute_summary(enriched)
    print("\n" + "=" * 70)
    print("AGGREGATE METRICS")
    print("=" * 70)
    for k, v in summary.items():
        print(f"{k}: {v:.4f}" if isinstance(v, float) else f"{k}: {v}")

    # 5. calibration
    cal = calibration_table(enriched, by="conf")
    cal.to_csv(output_dir / "04_calibration.csv", index=False)
    print("\n" + "=" * 70)
    print("CALIBRATION BY CONF BIN  (well-calibrated if conf ≈ hit_rate*100)")
    print("=" * 70)
    print(cal.to_string(index=False))
    print(f"\nAll outputs saved to {output_dir}/")
    return enriched, summary, cal


if __name__ == "__main__":
    # Allow CLI override: python backtest.py /path/to/reports
    reports = Path(sys.argv[1]) if len(sys.argv) > 1 else REPORTS_DIR
    if not reports.exists():
        print(f"[ERROR] Reports dir not found: {reports}")
        print("Edit REPORTS_DIR at top of file, or pass as argument:")
        print("  python backtest.py D:/path/to/reports")
        sys.exit(1)
    run(reports, price_fn=synthetic_price_fn)
